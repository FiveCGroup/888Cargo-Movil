# Generador de Documentación Backend 888Cargo con IA
# Integra OpenAI GPT para mejorar la calidad y detalle de la documentación
# Versión 2.0 - Con análisis inteligente de código y generación de contenido

import os
import sys
from datetime import datetime
from pathlib import Path
import json
import sqlite3
import asyncio
import aiohttp
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import ast
import re
from typing import Dict, List, Optional, Any
import hashlib
from dataclasses import dataclass
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich import print as rprint
import tiktoken

# Configuración de IA
@dataclass
class AIConfig:
    """Configuración para la integración de IA"""
    provider: str = "openai"  # "openai", "anthropic", "gemini"
    model: str = "gpt-4"
    api_key: str = ""
    max_tokens: int = 4000
    temperature: float = 0.3
    timeout: int = 30
    cache_enabled: bool = True
    cache_duration_hours: int = 24

class AIDocumentationEnhancer:
    """
    Mejorador de documentación con IA
    Analiza código y genera contenido técnico detallado
    """
    
    def __init__(self, config: AIConfig):
        self.config = config
        self.console = Console()
        self.encoding = tiktoken.encoding_for_model("gpt-4")
        self.cache = {}
        self.cache_file = Path("ai_cache.json")
        self.load_cache()
        
    def load_cache(self):
        """Carga cache de respuestas de IA"""
        if self.cache_file.exists():
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self.cache = json.load(f)
            except Exception as e:
                self.console.print(f"⚠️ Error cargando cache: {e}", style="yellow")
                self.cache = {}
                
    def save_cache(self):
        """Guarda cache de respuestas de IA"""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, indent=2, ensure_ascii=False)
        except Exception as e:
            self.console.print(f"⚠️ Error guardando cache: {e}", style="yellow")
            
    def get_cache_key(self, prompt: str, context: str = "") -> str:
        """Genera clave única para cache"""
        content = f"{prompt}|||{context}"
        return hashlib.md5(content.encode()).hexdigest()
        
    async def call_openai_api(self, messages: List[Dict], **kwargs) -> str:
        """Llama a la API de OpenAI"""
        headers = {
            "Authorization": f"Bearer {self.config.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.config.model,
            "messages": messages,
            "max_tokens": self.config.max_tokens,
            "temperature": self.config.temperature,
            **kwargs
        }
        
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=self.config.timeout)) as session:
            async with session.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=payload
            ) as response:
                if response.status == 200:
                    result = await response.json()
                    return result["choices"][0]["message"]["content"]
                else:
                    error_text = await response.text()
                    raise Exception(f"API Error {response.status}: {error_text}")
                    
    async def enhance_content(self, content_type: str, raw_data: Dict, context: str = "") -> str:
        """
        Mejora contenido usando IA
        
        Args:
            content_type: Tipo de contenido (function_analysis, architecture_overview, etc.)
            raw_data: Datos en bruto del análisis
            context: Contexto adicional
        """
        
        # Verificar cache
        cache_key = self.get_cache_key(str(raw_data), context)
        if self.config.cache_enabled and cache_key in self.cache:
            cache_entry = self.cache[cache_key]
            if self._is_cache_valid(cache_entry):
                return cache_entry["content"]
                
        # Preparar prompt según el tipo de contenido
        prompt = self._get_prompt_for_content_type(content_type, raw_data, context)
        
        try:
            messages = [
                {
                    "role": "system",
                    "content": """Eres un experto en documentación técnica de software. Tu trabajo es analizar código fuente y generar documentación detallada, clara y profesional en español. 

Características de tu escritura:
- Técnicamente precisa pero accesible
- Incluye ejemplos prácticos cuando sea relevante
- Explica el "por qué" además del "qué"
- Identifica patrones de diseño y buenas prácticas
- Señala posibles mejoras o consideraciones
- Usa un tono profesional pero no demasiado formal
- Incluye emojis apropiados para mejorar la legibilidad

Formato de respuesta:
- Usa markdown para estructurar el contenido
- Incluye código cuando sea necesario
- Organiza la información de manera lógica
- No repitas información obvia"""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
            
            enhanced_content = await self.call_openai_api(messages)
            
            # Guardar en cache
            if self.config.cache_enabled:
                self.cache[cache_key] = {
                    "content": enhanced_content,
                    "timestamp": datetime.now().isoformat(),
                    "content_type": content_type
                }
                self.save_cache()
                
            return enhanced_content
            
        except Exception as e:
            self.console.print(f"❌ Error en IA para {content_type}: {e}", style="red")
            return self._get_fallback_content(content_type, raw_data)
            
    def _is_cache_valid(self, cache_entry: Dict) -> bool:
        """Verifica si una entrada de cache sigue siendo válida"""
        try:
            timestamp = datetime.fromisoformat(cache_entry["timestamp"])
            age_hours = (datetime.now() - timestamp).total_seconds() / 3600
            return age_hours < self.config.cache_duration_hours
        except:
            return False
            
    def _get_prompt_for_content_type(self, content_type: str, raw_data: Dict, context: str) -> str:
        """Genera prompt específico según el tipo de contenido"""
        
        prompts = {
            "function_analysis": f"""
Analiza esta función JavaScript/Node.js y proporciona documentación detallada:

**Archivo:** {raw_data.get('file_path', 'N/A')}
**Funciones encontradas:** {raw_data.get('functions', [])}
**Líneas de código:** {raw_data.get('lines', 0)}
**Imports:** {raw_data.get('imports', [])}
**Exports:** {raw_data.get('exports', [])}

**Código fuente (muestra):**
```javascript
{raw_data.get('code_sample', 'No disponible')}
```

**Contexto del proyecto:** {context}

Por favor, proporciona:
1. **Propósito y responsabilidad** del archivo
2. **Análisis de las funciones principales** (qué hacen, parámetros, retorno)
3. **Patrones de diseño** utilizados
4. **Dependencias** y su propósito
5. **Posibles mejoras** o consideraciones
6. **Ejemplos de uso** si es relevante

Responde en español, con formato markdown y enfoque en la utilidad para desarrolladores.
""",

            "architecture_overview": f"""
Analiza esta arquitectura de software y proporciona una visión general:

**Estructura del proyecto:**
{json.dumps(raw_data.get('directory_structure', {}), indent=2)}

**Tecnologías identificadas:**
{raw_data.get('technologies', [])}

**Patrones detectados:**
{raw_data.get('patterns', [])}

**Base de datos:**
{json.dumps(raw_data.get('database_schema', {}), indent=2)}

Por favor, proporciona:
1. **Visión general** de la arquitectura
2. **Patrones de diseño** implementados y su justificación
3. **Flujo de datos** principal
4. **Separación de responsabilidades**
5. **Escalabilidad** y consideraciones de rendimiento
6. **Buenas prácticas** observadas
7. **Recomendaciones** de mejora

Responde en español, con formato técnico profesional.
""",

            "api_documentation": f"""
Documenta estos endpoints de API:

**Rutas encontradas:**
{json.dumps(raw_data.get('routes', []), indent=2)}

**Middlewares aplicados:**
{raw_data.get('middlewares', [])}

**Controladores:**
{raw_data.get('controllers', [])}

**Modelos de datos:**
{raw_data.get('models', [])}

Por favor, proporciona:
1. **Documentación de cada endpoint** (método, ruta, propósito)
2. **Parámetros de entrada** esperados
3. **Respuestas típicas** y códigos de estado
4. **Autenticación** requerida
5. **Ejemplos de uso** con curl o JavaScript
6. **Manejo de errores**
7. **Rate limiting** y consideraciones de seguridad

Usa formato de documentación API estándar.
""",

            "database_analysis": f"""
Analiza este esquema de base de datos:

**Tablas identificadas:**
{json.dumps(raw_data.get('tables', {}), indent=2)}

**Relaciones:**
{raw_data.get('relationships', [])}

**Índices:**
{raw_data.get('indexes', [])}

Por favor, proporciona:
1. **Descripción del modelo de datos**
2. **Propósito de cada tabla**
3. **Relaciones entre entidades**
4. **Índices y optimizaciones**
5. **Integridad de datos** y constraints
6. **Consideraciones de rendimiento**
7. **Posibles mejoras** al esquema

Enfócate en la lógica de negocio y el diseño de datos.
""",

            "security_analysis": f"""
Analiza los aspectos de seguridad del sistema:

**Middlewares de seguridad:**
{raw_data.get('security_middlewares', [])}

**Autenticación:**
{raw_data.get('auth_methods', [])}

**Validaciones:**
{raw_data.get('validations', [])}

**Rate limiting:**
{raw_data.get('rate_limiting', {})}

Por favor, proporciona:
1. **Análisis de seguridad** general
2. **Autenticación y autorización**
3. **Validación de datos** de entrada
4. **Prevención de ataques** comunes (XSS, SQL injection, etc.)
5. **Rate limiting** y DoS protection
6. **Buenas prácticas** implementadas
7. **Recomendaciones** de mejora de seguridad

Enfócate en aspectos prácticos de seguridad web.
"""
        }
        
        return prompts.get(content_type, f"Analiza y documenta: {raw_data}")
        
    def _get_fallback_content(self, content_type: str, raw_data: Dict) -> str:
        """Contenido de respaldo cuando la IA no está disponible"""
        fallbacks = {
            "function_analysis": f"""
## Análisis de Archivo

**Ubicación:** {raw_data.get('file_path', 'N/A')}
**Líneas de código:** {raw_data.get('lines', 0)}
**Funciones encontradas:** {len(raw_data.get('functions', []))}

### Funciones Principales
{chr(10).join([f"- {func}" for func in raw_data.get('functions', [])])}

### Dependencias
{chr(10).join([f"- {imp}" for imp in raw_data.get('imports', [])])}
""",
            "architecture_overview": """
## Visión General de la Arquitectura

El sistema implementa una arquitectura en capas con separación clara de responsabilidades:

- **Capa de Presentación:** Manejo de requests HTTP
- **Capa de Aplicación:** Lógica de negocio y servicios
- **Capa de Datos:** Persistencia y acceso a datos
""",
            "api_documentation": """
## Documentación de API

### Endpoints Principales
- Autenticación y gestión de usuarios
- Gestión de cargas y packing lists  
- Generación y validación de códigos QR

### Formatos de Respuesta
Todas las respuestas siguen el formato JSON estándar con códigos de estado HTTP apropiados.
""",
        }
        
        return fallbacks.get(content_type, "## Contenido no disponible\n\nNo se pudo generar contenido mejorado para esta sección.")

class EnhancedBackendDocumentationGenerator:
    """
    Generador mejorado de documentación con IA
    Versión 2.0 con análisis inteligente y contenido mejorado
    """
    
    def __init__(self, backend_path, output_path, ai_config: AIConfig):
        self.backend_path = Path(backend_path)
        self.output_path = Path(output_path)
        self.ai_enhancer = AIDocumentationEnhancer(ai_config)
        self.console = Console()
        self.doc = Document()
        self.current_date = datetime.now().strftime("%d de %B de %Y")
        
        # Configuración de estilos
        self.setup_styles()
        
        # Datos del proyecto
        self.project_info = {
            'name': '888Cargo Backend',
            'version': '1.0.0', 
            'description': 'Sistema de gestión de listas de empaque con códigos QR',
            'author': 'FiveCGroup',
            'license': 'MIT'
        }
        
        # Estructura de directorios a analizar
        self.directories_to_analyze = [
            'controllers', 'services', 'models', 'repositories',
            'routes', 'middlewares', 'validators', 'utils', 'config'
        ]
        
        # Datos recopilados para análisis de IA
        self.collected_data = {
            'files_analyzed': [],
            'technologies': set(),
            'patterns': set(),
            'security_features': [],
            'api_endpoints': [],
            'database_tables': {}
        }
        
    def setup_styles(self):
        """Configura estilos mejorados para el documento"""
        
        # Estilo para título principal con gradiente simulado
        title_style = self.doc.styles.add_style('EnhancedTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Segoe UI'
        title_font.size = Pt(32)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 102, 204)  # Azul corporativo
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Estilo para subtítulos con iconos
        h1_style = self.doc.styles.add_style('EnhancedH1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Segoe UI Semibold'
        h1_font.size = Pt(22)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 102, 204)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(16)
        h1_style.paragraph_format.keep_with_next = True
        
        # Estilo para código mejorado
        code_style = self.doc.styles.add_style('EnhancedCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'JetBrains Mono'
        code_font.size = Pt(9)
        code_font.color.rgb = RGBColor(0, 0, 0)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Estilo para destacados
        highlight_style = self.doc.styles.add_style('Highlight', WD_STYLE_TYPE.PARAGRAPH)
        highlight_font = highlight_style.font
        highlight_font.name = 'Segoe UI'
        highlight_font.size = Pt(11)
        highlight_font.italic = True
        highlight_font.color.rgb = RGBColor(51, 102, 153)
        highlight_style.paragraph_format.left_indent = Inches(0.3)
        highlight_style.paragraph_format.space_before = Pt(6)
        highlight_style.paragraph_format.space_after = Pt(6)
        
    async def analyze_file_with_ai(self, file_path: Path) -> Dict:
        """Analiza un archivo usando IA para obtener insights detallados"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            self.console.print(f"⚠️ Error leyendo {file_path}: {e}", style="yellow")
            return {}
            
        # Análisis básico del código
        basic_analysis = self._basic_code_analysis(content, file_path)
        
        # Preparar datos para IA
        ai_data = {
            'file_path': str(file_path.relative_to(self.backend_path)),
            'code_sample': content[:2000] if len(content) > 2000 else content,  # Limitar para tokens
            **basic_analysis
        }
        
        # Contexto del proyecto
        context = f"Proyecto: {self.project_info['name']} - {self.project_info['description']}"
        
        # Obtener análisis mejorado de IA
        enhanced_analysis = await self.ai_enhancer.enhance_content(
            "function_analysis", 
            ai_data, 
            context
        )
        
        return {
            'basic': basic_analysis,
            'enhanced': enhanced_analysis,
            'file_path': file_path
        }
        
    def _basic_code_analysis(self, content: str, file_path: Path) -> Dict:
        """Análisis básico del código sin IA"""
        
        analysis = {
            'lines': len(content.split('\n')),
            'functions': [],
            'classes': [],
            'exports': [],
            'imports': [],
            'comments': [],
            'complexity_score': 0
        }
        
        # Buscar funciones con mejor regex
        function_patterns = [
            r'(?:async\s+)?function\s+(\w+)',  # function declarations
            r'const\s+(\w+)\s*=\s*(?:async\s+)?(?:function|\([^)]*\)\s*=>)',  # arrow functions
            r'(\w+)\s*:\s*(?:async\s+)?(?:function|\([^)]*\)\s*=>)',  # object methods
            r'static\s+(?:async\s+)?(\w+)\s*\(',  # static methods
        ]
        
        for pattern in function_patterns:
            for match in re.finditer(pattern, content):
                func_name = match.group(1)
                if func_name and func_name not in analysis['functions']:
                    analysis['functions'].append(func_name)
                    
        # Buscar clases
        class_pattern = r'class\s+(\w+)(?:\s+extends\s+\w+)?'
        for match in re.finditer(class_pattern, content):
            analysis['classes'].append(match.group(1))
            
        # Buscar exports más completos
        export_patterns = [
            r'module\.exports\s*=\s*(\w+)',
            r'export\s+(?:default\s+)?(?:class\s+|function\s+|const\s+)?(\w+)',
            r'exports\.(\w+)\s*='
        ]
        
        for pattern in export_patterns:
            for match in re.finditer(pattern, content):
                if match.group(1):
                    analysis['exports'].append(match.group(1))
                    
        # Buscar imports mejorados
        import_patterns = [
            r'require\([\'"]([^\'"]+)[\'"]\)',
            r'import.*from\s+[\'"]([^\'"]+)[\'"]',
            r'import\s+[\'"]([^\'"]+)[\'"]'
        ]
        
        for pattern in import_patterns:
            for match in re.finditer(pattern, content):
                module = match.group(1)
                if module and not module.startswith('.'):
                    analysis['imports'].append(module)
                    
        # Calcular complejidad ciclomática básica
        complexity_keywords = ['if', 'else', 'while', 'for', 'switch', 'case', 'catch', '&&', '||', '?']
        for keyword in complexity_keywords:
            analysis['complexity_score'] += len(re.findall(rf'\b{keyword}\b', content))
            
        # Buscar comentarios importantes
        comment_patterns = [
            r'//\s*(.{20,})',  # Comentarios de línea largos
            r'/\*\*(.*?)\*/',  # JSDoc comments
            r'//\s*TODO:?\s*(.+)',  # TODOs
            r'//\s*FIXME:?\s*(.+)',  # FIXMEs
        ]
        
        for pattern in comment_patterns:
            for match in re.finditer(pattern, content, re.DOTALL):
                comment = match.group(1).strip()
                if len(comment) > 10:
                    analysis['comments'].append(comment[:100])  # Limitar longitud
                    
        return analysis
        
    async def generate_enhanced_documentation(self):
        """Genera documentación completa mejorada con IA"""
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=self.console,
        ) as progress:
            
            # Task para progreso general
            main_task = progress.add_task("🚀 Generando documentación...", total=100)
            
            # Análisis inicial
            progress.update(main_task, advance=5, description="📊 Analizando estructura del proyecto...")
            await asyncio.sleep(0.1)  # Para mostrar progreso
            
            self.analyze_package_json()
            self.analyze_project_structure()
            
            # Crear documento base
            progress.update(main_task, advance=10, description="📄 Creando documento base...")
            self.add_enhanced_title_page()
            self.add_enhanced_table_of_contents()
            
            # Análisis de archivos con IA
            progress.update(main_task, advance=5, description="🤖 Analizando archivos con IA...")
            
            for i, directory in enumerate(self.directories_to_analyze):
                dir_path = self.backend_path / directory
                if dir_path.exists():
                    task_desc = f"🔍 Analizando {directory}/ con IA..."
                    progress.update(main_task, description=task_desc)
                    
                    await self.analyze_directory_with_ai(directory, dir_path)
                    
                    # Calcular progreso (20-70% para análisis de directorios)
                    dir_progress = 20 + (i + 1) * (50 / len(self.directories_to_analyze))
                    progress.update(main_task, completed=dir_progress)
                    
            # Análisis de base de datos
            progress.update(main_task, advance=10, description="🗄️ Analizando base de datos...")
            db_schema = self.analyze_database_schema()
            if db_schema:
                await self.generate_enhanced_database_section(db_schema)
                
            # Generar secciones con IA
            progress.update(main_task, advance=5, description="🏗️ Generando arquitectura con IA...")
            await self.generate_enhanced_architecture_section()
            
            progress.update(main_task, advance=5, description="🔒 Analizando seguridad con IA...")
            await self.generate_enhanced_security_section()
            
            progress.update(main_task, advance=5, description="📡 Documentando API con IA...")
            await self.generate_enhanced_api_section()
            
            # Finalizar documento
            progress.update(main_task, advance=5, description="✨ Finalizando documento...")
            self.add_enhanced_conclusions()
            
            # Guardar
            output_file = self.output_path / f"888Cargo_Backend_Documentation_AI_Enhanced_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            self.doc.save(output_file)
            
            progress.update(main_task, completed=100, description="✅ ¡Documentación completada!")
            
            return output_file
            
    def add_enhanced_title_page(self):
        """Crea una página de título mejorada"""
        
        # Título principal con estilo mejorado
        title = self.doc.add_paragraph('📋 DOCUMENTACIÓN TÉCNICA COMPLETA', style='EnhancedTitle')
        
        # Subtítulo del proyecto
        subtitle = self.doc.add_paragraph('🚀 888CARGO BACKEND SYSTEM', style='EnhancedH1')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Banner de descripción
        banner = self.doc.add_paragraph(
            '🎯 Sistema Avanzado de Gestión de Listas de Empaque\n'
            'con Generación Automática de Códigos QR y Autenticación JWT'
        )
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner.runs[0].font.size = Pt(14)
        banner.runs[0].font.italic = True
        
        # Espacio
        self.doc.add_paragraph('\n')
        
        # Información del documento en tabla mejorada
        info_table = self.doc.add_table(rows=8, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ['📅 Fecha de Generación', self.current_date],
            ['🚀 Versión del Sistema', self.project_info['version']],
            ['🤖 Generado con IA', 'OpenAI GPT-4'],
            ['👨‍💻 Desarrollado por', self.project_info['author']],
            ['📄 Licencia', self.project_info['license']],
            ['⚙️ Tecnologías Principales', 'Node.js + Express.js + SQLite'],
            ['🔒 Nivel de Seguridad', 'Empresarial - JWT + Validaciones'],
            ['📊 Estado del Sistema', 'Producción - Estable y Escalable']
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Estilo para las celdas
            for j, cell in enumerate(row.cells):
                cell.paragraphs[0].runs[0].font.name = 'Segoe UI'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                if j == 0:  # Primera columna en negrita
                    cell.paragraphs[0].runs[0].font.bold = True
                    
        # Nota de IA
        ai_note = self.doc.add_paragraph(
            '\n🤖 DOCUMENTACIÓN MEJORADA CON INTELIGENCIA ARTIFICIAL\n'
            'Este documento ha sido generado automáticamente con análisis inteligente de código, '
            'arquitectura y mejores prácticas. El contenido incluye insights generados por IA '
            'para proporcionar documentación técnica de nivel profesional.'
        )
        ai_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ai_note.runs[0].font.italic = True
        ai_note.runs[0].font.size = Pt(10)
        ai_note.runs[0].font.color.rgb = RGBColor(0, 102, 153)
        
        self.add_page_break()
        
    async def analyze_directory_with_ai(self, dir_name: str, dir_path: Path):
        """Analiza un directorio completo usando IA"""
        
        self.add_page_break()
        
        # Título de sección con emoji
        section_emojis = {
            'controllers': '🎮',
            'services': '⚙️',
            'models': '📊',
            'repositories': '🗄️',
            'routes': '🛣️',
            'middlewares': '🔒',
            'validators': '✅',
            'utils': '🛠️',
            'config': '⚙️'
        }
        
        emoji = section_emojis.get(dir_name, '📁')
        section_title = f'{emoji} {dir_name.upper()} - ANÁLISIS INTELIGENTE'
        self.doc.add_paragraph(section_title, style='EnhancedH1')
        
        # Descripción mejorada del directorio
        descriptions = {
            'controllers': 'Los controladores actúan como orquestadores principales, coordinando requests HTTP con servicios de negocio.',
            'services': 'Los servicios encapsulan la lógica de negocio compleja y proporcionan APIs internas reutilizables.',
            'models': 'Los modelos definen la estructura de datos del dominio y las reglas de validación empresarial.',
            'repositories': 'Los repositorios abstraen completamente el acceso a datos, implementando el patrón Repository.',
            'routes': 'Las rutas definen la API REST pública y aplican middlewares de seguridad y validación.',
            'middlewares': 'Los middlewares implementan funcionalidades transversales críticas para la seguridad y monitoreo.',
            'validators': 'Los validadores aseguran la integridad de datos mediante esquemas de validación robustos.',
            'utils': 'Las utilidades proporcionan funciones auxiliares especializadas y reutilizables.',
            'config': 'Las configuraciones centralizan parámetros del sistema y gestión de entornos.'
        }
        
        self.doc.add_paragraph(descriptions.get(dir_name, f'Componentes del directorio {dir_name}.'))
        
        # Analizar archivos JavaScript en el directorio
        js_files = list(dir_path.glob('*.js'))
        
        if js_files:
            # Información general
            summary_para = self.doc.add_paragraph(f'📊 **Resumen:** {len(js_files)} archivos encontrados')
            summary_para.runs[0].font.bold = True
            
            # Analizar cada archivo con IA
            for js_file in js_files[:5]:  # Limitar a 5 archivos por directorio para tokens
                try:
                    analysis = await self.analyze_file_with_ai(js_file)
                    if analysis:
                        await self.add_enhanced_file_analysis(analysis)
                        
                        # Recopilar datos para análisis global
                        self.collected_data['files_analyzed'].append(analysis)
                        
                        if 'basic' in analysis:
                            self.collected_data['technologies'].update(analysis['basic'].get('imports', []))
                            
                except Exception as e:
                    self.console.print(f"⚠️ Error analizando {js_file}: {e}", style="yellow")
                    
    async def add_enhanced_file_analysis(self, analysis: Dict):
        """Añade análisis mejorado de archivo al documento"""
        
        if not analysis:
            return
            
        basic = analysis.get('basic', {})
        enhanced = analysis.get('enhanced', '')
        file_path = analysis.get('file_path')
        
        # Nombre del archivo con indicadores
        file_paragraph = self.doc.add_paragraph()
        
        # Icono según tipo de archivo
        file_icons = {
            'controller': '🎮',
            'service': '⚙️', 
            'model': '📊',
            'repository': '🗄️',
            'route': '🛣️',
            'middleware': '🔒',
            'validator': '✅',
            'util': '🛠️',
            'config': '⚙️'
        }
        
        icon = '📄'
        for key, emoji in file_icons.items():
            if key in str(file_path).lower():
                icon = emoji
                break
                
        run = file_paragraph.add_run(f'{icon} {file_path}')
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Métricas en tabla compacta
        if basic:
            metrics_table = self.doc.add_table(rows=1, cols=4)
            metrics_table.style = 'Light List Accent 1'
            
            headers = ['📏 Líneas', '🔧 Funciones', '📦 Imports', '🎯 Complejidad']
            values = [
                str(basic.get('lines', 0)),
                str(len(basic.get('functions', []))),
                str(len(basic.get('imports', []))),
                str(basic.get('complexity_score', 0))
            ]
            
            for i, (header, value) in enumerate(zip(headers, values)):
                cell = metrics_table.rows[0].cells[i]
                cell.text = f'{header}\n{value}'
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        # Contenido mejorado por IA
        if enhanced:
            self.doc.add_paragraph(enhanced)
        else:
            # Fallback básico
            if basic.get('functions'):
                func_para = self.doc.add_paragraph('🔧 **Funciones principales:**')
                func_para.runs[0].font.bold = True
                for func in basic['functions'][:5]:
                    self.doc.add_paragraph(f'  • {func}')
                    
        self.doc.add_paragraph('')  # Espacio entre archivos
        
    def add_page_break(self):
        """Añade salto de página"""
        self.doc.add_page_break()
        
    # Métodos adicionales necesarios...
    def analyze_package_json(self):
        """Analiza package.json (implementación básica)"""
        package_path = self.backend_path.parent / 'package.json'
        if package_path.exists():
            try:
                with open(package_path, 'r', encoding='utf-8') as f:
                    package_data = json.load(f)
                    self.project_info.update({
                        'name': package_data.get('name', self.project_info['name']),
                        'version': package_data.get('version', self.project_info['version']),
                        'dependencies': package_data.get('dependencies', {}),
                        'scripts': package_data.get('scripts', {})
                    })
            except Exception as e:
                self.console.print(f"⚠️ Error analizando package.json: {e}", style="yellow")
                
    def analyze_project_structure(self):
        """Analiza estructura general del proyecto"""
        # Implementación básica
        pass
        
    def analyze_database_schema(self):
        """Analiza esquema de base de datos"""
        # Implementación básica (reutilizar del generador anterior)
        return {}
        
    async def generate_enhanced_architecture_section(self):
        """Genera sección de arquitectura mejorada"""
        try:
            self.console.print("🏗️ Generando análisis de arquitectura...", style="cyan")
            
            # Obtener información de la estructura del proyecto
            project_structure = self._analyze_project_structure()
            
            architecture_prompt = f"""
            Analiza la arquitectura del siguiente proyecto backend Node.js:
            
            Estructura del proyecto:
            {json.dumps(project_structure, indent=2, ensure_ascii=False)}
            
            Genera un análisis detallado que incluya:
            1. Patrón arquitectónico identificado
            2. Organización de carpetas y módulos
            3. Separación de responsabilidades
            4. Flujo de datos entre capas
            5. Patrones de diseño implementados
            6. Escalabilidad y mantenibilidad
            7. Recomendaciones de mejora
            
            Formato: Análisis técnico profesional en español.
            """
            
            response = await self._make_ai_request(
                architecture_prompt,
                cache_key="architecture_analysis"
            )
            
            # Añadir sección al documento
            self.doc.add_paragraph('🏗️ ANÁLISIS DE ARQUITECTURA', style='EnhancedH1')
            self.doc.add_paragraph(response or self._get_fallback_architecture_content())
            
        except Exception as e:
            self.console.print(f"[red]Error generando sección de arquitectura: {e}")
            self.doc.add_paragraph('🏗️ ANÁLISIS DE ARQUITECTURA', style='EnhancedH1')
            self.doc.add_paragraph(self._get_fallback_architecture_content())
    
    def _get_fallback_architecture_content(self):
        """Contenido de arquitectura de respaldo"""
        return """
        ## Análisis de Arquitectura

        ### Patrón Arquitectónico
        El proyecto implementa una arquitectura en capas (Layered Architecture) con separación clara de responsabilidades:
        
        - **Capa de Presentación**: Routes y Controllers
        - **Capa de Lógica de Negocio**: Services
        - **Capa de Datos**: Models y Repositories
        - **Capa de Infraestructura**: Database y Middlewares

        ### Organización del Código
        - /controllers: Manejo de requests HTTP
        - /services: Lógica de negocio
        - /models: Definición de entidades
        - /routes: Definición de endpoints
        - /middlewares: Funciones intermedias
        - /utils: Utilidades generales

        ### Ventajas de la Arquitectura
        - Separación clara de responsabilidades
        - Fácil mantenimiento y testing
        - Escalabilidad horizontal
        - Reutilización de componentes

        ### Recomendaciones
        - Implementar inyección de dependencias
        - Añadir layer de validación
        - Considerar patrones CQRS para operaciones complejas
        """
        
    async def generate_enhanced_security_section(self):
        """Genera sección de seguridad mejorada"""
        try:
            self.console.print("🔒 Generando análisis de seguridad...", style="cyan")
            
            # Obtener información de endpoints
            endpoints_info = self._extract_endpoints_info()
            
            security_prompt = f"""
            Analiza los siguientes endpoints de una API REST desde una perspectiva de seguridad:
            
            {json.dumps(endpoints_info, indent=2, ensure_ascii=False)}
            
            Proporciona un análisis detallado que incluya:
            1. Vulnerabilidades potenciales identificadas
            2. Recomendaciones específicas para cada endpoint
            3. Mejores prácticas de seguridad aplicables
            4. Validaciones de entrada necesarias
            5. Manejo de autenticación y autorización
            6. Protección contra ataques comunes (SQL injection, XSS, CSRF)
            
            Formato: Texto estructurado en español, profesional y técnico.
            """
            
            response = await self._make_ai_request(
                security_prompt,
                cache_key="security_analysis"
            )
            
            # Añadir sección al documento
            self.doc.add_paragraph('🔒 ANÁLISIS DE SEGURIDAD', style='EnhancedH1')
            self.doc.add_paragraph(response or self._get_fallback_security_content())
            
        except Exception as e:
            self.console.print(f"[red]Error generando sección de seguridad: {e}")
            self.doc.add_paragraph('🔒 ANÁLISIS DE SEGURIDAD', style='EnhancedH1')
            self.doc.add_paragraph(self._get_fallback_security_content())
    
    def _get_fallback_security_content(self):
        """Contenido de seguridad de respaldo"""
        return """
        ## Consideraciones de Seguridad

        ### Autenticación y Autorización
        - Implementar JWT tokens para autenticación
        - Validar permisos de usuario en cada endpoint
        - Usar middlewares de autenticación apropiados

        ### Validación de Datos
        - Validar todos los parámetros de entrada
        - Sanitizar datos antes del procesamiento
        - Implementar límites de tasa (rate limiting)

        ### Protección contra Ataques
        - SQL Injection: Usar consultas preparadas
        - XSS: Sanitizar salidas HTML
        - CSRF: Implementar tokens CSRF
        - CORS: Configurar políticas apropiadas

        ### Recomendaciones Adicionales
        - Usar HTTPS en producción
        - Implementar logging de seguridad
        - Monitorear actividad sospechosa
        - Mantener dependencias actualizadas
        """
        
    async def generate_enhanced_api_section(self):
        """Genera sección de API mejorada"""
        try:
            self.console.print("📡 Generando documentación de API mejorada...", style="cyan")
            
            # Obtener información detallada de la API
            api_info = self._extract_comprehensive_api_info()
            
            api_prompt = f"""
            Genera documentación técnica detallada para esta API REST:
            
            {json.dumps(api_info, indent=2, ensure_ascii=False)}
            
            Incluye:
            1. Descripción general de la arquitectura API
            2. Patrones de diseño utilizados
            3. Documentación detallada de cada endpoint
            4. Códigos de respuesta HTTP
            5. Ejemplos de requests y responses
            6. Mejores prácticas implementadas
            7. Consideraciones de performance
            
            Formato: Documentación técnica profesional en español.
            """
            
            response = await self._make_ai_request(
                api_prompt,
                cache_key="api_documentation"
            )
            
            # Añadir sección al documento
            self.doc.add_paragraph('📡 DOCUMENTACIÓN DE API', style='EnhancedH1')
            self.doc.add_paragraph(response or self._get_fallback_api_content())
            
        except Exception as e:
            self.console.print(f"[red]Error generando sección de API: {e}")
            self.doc.add_paragraph('📡 DOCUMENTACIÓN DE API', style='EnhancedH1')
            self.doc.add_paragraph(self._get_fallback_api_content())
    
    def _get_fallback_api_content(self):
        """Contenido de API de respaldo"""
        return """
        ## Documentación de API REST

        ### Arquitectura
        La API sigue los principios REST con endpoints organizados por recursos.
        Utiliza JSON para el intercambio de datos y códigos de estado HTTP estándar.

        ### Endpoints Principales
        - GET /api/cargas - Listar cargas
        - POST /api/cargas - Crear nueva carga
        - PUT /api/cargas/:id - Actualizar carga
        - DELETE /api/cargas/:id - Eliminar carga

        ### Códigos de Respuesta
        - 200: Operación exitosa
        - 201: Recurso creado
        - 400: Error en la solicitud
        - 401: No autorizado
        - 404: Recurso no encontrado
        - 500: Error interno del servidor

        ### Formato de Respuestas
        Todas las respuestas incluyen:
        - success: boolean
        - data: objeto/array de datos
        - message: mensaje descriptivo
        - timestamp: marca de tiempo
        """
        
    def add_enhanced_table_of_contents(self):
        """Tabla de contenidos mejorada"""
        # Implementación básica
        self.doc.add_paragraph('📋 TABLA DE CONTENIDOS', style='EnhancedH1')
        
    def _extract_endpoints_info(self):
        """Extrae información de endpoints para análisis de seguridad"""
        endpoints = []
        try:
            routes_dir = Path("routes")
            if routes_dir.exists():
                for route_file in routes_dir.glob("*.js"):
                    with open(route_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                        
                        # Buscar definiciones de rutas
                        route_patterns = [
                            r'router\.get\([\'"]([^\'"]+)[\'"]',
                            r'router\.post\([\'"]([^\'"]+)[\'"]',
                            r'router\.put\([\'"]([^\'"]+)[\'"]',
                            r'router\.delete\([\'"]([^\'"]+)[\'"]'
                        ]
                        
                        for pattern in route_patterns:
                            matches = re.findall(pattern, content)
                            for match in matches:
                                endpoints.append({
                                    'file': route_file.name,
                                    'method': pattern.split('\\')[1].replace('.', '').upper(),
                                    'path': match
                                })
        except Exception as e:
            self.console.print(f"[yellow]Warning: Error extrayendo endpoints: {e}")
        
        return endpoints
    
    def _extract_comprehensive_api_info(self):
        """Extrae información completa de la API"""
        api_info = {
            'endpoints': self._extract_endpoints_info(),
            'middlewares': [],
            'models': [],
            'controllers': []
        }
        
        try:
            # Analizar middlewares
            middleware_dir = Path("middlewares")
            if middleware_dir.exists():
                api_info['middlewares'] = [f.name for f in middleware_dir.glob("*.js")]
            
            # Analizar modelos
            models_dir = Path("models")
            if models_dir.exists():
                api_info['models'] = [f.name for f in models_dir.glob("*.js")]
            
            # Analizar controladores
            controllers_dir = Path("controllers")
            if controllers_dir.exists():
                api_info['controllers'] = [f.name for f in controllers_dir.glob("*.js")]
                
        except Exception as e:
            self.console.print(f"[yellow]Warning: Error analizando API: {e}")
        
        return api_info
    
    def _analyze_project_structure(self):
        """Analiza la estructura del proyecto"""
        structure = {}
        
        try:
            # Analizar directorios principales
            main_dirs = ['controllers', 'models', 'routes', 'services', 'middlewares', 'utils', 'config']
            
            for dir_name in main_dirs:
                dir_path = Path(dir_name)
                if dir_path.exists():
                    files = list(dir_path.glob("*.js"))
                    structure[dir_name] = {
                        'files_count': len(files),
                        'files': [f.name for f in files[:10]]  # Limitar a 10 archivos
                    }
            
            # Analizar package.json
            if Path("package.json").exists():
                with open("package.json", 'r', encoding='utf-8') as f:
                    package_data = json.load(f)
                    structure['dependencies'] = list(package_data.get('dependencies', {}).keys())[:15]
                    
        except Exception as e:
            self.console.print(f"[yellow]Warning: Error analizando estructura: {e}")
        
        return structure
    
    def add_enhanced_conclusions(self):
        """Añade conclusiones y recomendaciones"""
        try:
            self.doc.add_paragraph('🎯 CONCLUSIONES Y RECOMENDACIONES', style='EnhancedH1')
            
            conclusions = """
            ## Evaluación General del Proyecto
            
            El backend de 888Cargo presenta una arquitectura sólida basada en Node.js y Express, 
            con una estructura bien organizada que facilita el mantenimiento y la escalabilidad.
            
            ### Fortalezas Identificadas
            ✅ Estructura modular clara y bien organizada
            ✅ Separación apropiada entre capas (routes, controllers, services)
            ✅ Implementación de middlewares de seguridad
            ✅ Manejo estructurado de errores
            ✅ Configuración flexible del entorno
            
            ### Áreas de Mejora
            🔧 Implementar más validaciones de datos de entrada
            🔧 Añadir más pruebas unitarias e integración
            🔧 Mejorar documentación de API (OpenAPI/Swagger)
            🔧 Implementar monitoreo y logging avanzado
            🔧 Considerar migración a TypeScript para mayor robustez
            
            ### Recomendaciones Técnicas
            
            #### Corto Plazo (1-2 meses)
            - Implementar validación de entrada más robusta
            - Añadir rate limiting a endpoints críticos
            - Configurar CI/CD pipeline
            - Implementar health checks
            
            #### Mediano Plazo (3-6 meses)
            - Migrar a TypeScript
            - Implementar caché distribuido (Redis)
            - Añadir métricas de performance
            - Configurar monitoreo con alertas
            
            #### Largo Plazo (6+ meses)
            - Considerar arquitectura de microservicios
            - Implementar Event Sourcing para auditoría
            - Añadir machine learning para optimizaciones
            - Explorar tecnologías de containerización
            
            ### Conclusión Final
            
            El proyecto 888Cargo Backend está bien encaminado con una base sólida que permite
            crecimiento y evolución. Las recomendaciones propuestas mejorarán la robustez,
            seguridad y mantenibilidad del sistema.
            """
            
            self.doc.add_paragraph(conclusions)
            
        except Exception as e:
            self.console.print(f"[red]Error añadiendo conclusiones: {e}")
            self.doc.add_paragraph('🎯 CONCLUSIONES', style='EnhancedH1')
            self.doc.add_paragraph("Sección de conclusiones en desarrollo.")

async def main():
    """Función principal mejorada"""
    console = Console()
    
    console.print("🚀 Generador de Documentación Backend 888Cargo con IA", style="bold green")
    console.print("=" * 60, style="blue")
    
    # Verificar API key de OpenAI
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        console.print("\n❌ Error: OPENAI_API_KEY no encontrada", style="red")
        console.print("📋 Para configurar:", style="yellow")
        console.print("   1. Obtén tu API key de: https://platform.openai.com/api-keys", style="white")
        console.print("   2. Ejecuta: $env:OPENAI_API_KEY='tu-api-key-aqui'", style="white")
        console.print("   3. O añádela al archivo .env", style="white")
        sys.exit(1)
        
    # Configurar rutas
    backend_path = Path(__file__).parent
    output_path = backend_path / 'docs'
    output_path.mkdir(exist_ok=True)
    
    # Configuración de IA
    ai_config = AIConfig(
        provider="openai",
        model="gpt-4",
        api_key=api_key,
        max_tokens=3000,
        temperature=0.3,
        cache_enabled=True
    )
    
    try:
        # Crear generador mejorado
        generator = EnhancedBackendDocumentationGenerator(backend_path, output_path, ai_config)
        
        # Generar documentación
        output_file = await generator.generate_enhanced_documentation()
        
        console.print("\n" + "=" * 60, style="blue")
        console.print("✅ ¡Documentación con IA generada exitosamente!", style="bold green")
        console.print(f"📁 Archivo: {output_file}", style="cyan")
        console.print(f"💾 Tamaño: {output_file.stat().st_size / 1024:.2f} KB", style="cyan")
        console.print("🤖 Mejorada con análisis de IA", style="magenta")
        
        # Abrir archivo automáticamente
        if os.name == 'nt':
            os.startfile(output_file)
            console.print("📖 Abriendo documento...", style="green")
            
    except Exception as e:
        console.print(f"❌ Error generando documentación: {e}", style="red")
        console.print("💡 Verifica tu API key y conexión a internet", style="yellow")
        sys.exit(1)

if __name__ == "__main__":
    asyncio.run(main())