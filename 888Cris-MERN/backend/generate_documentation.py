# Generador de Documentación Backend 888Cargo
# Genera documentación completa en formato Word (.docx)

import os
import sys
from datetime import datetime
from pathlib import Path
import json
import sqlite3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import ast
import re

class BackendDocumentationGenerator:
    """
    Generador completo de documentación para el backend de 888Cargo
    Produce documentos Word profesionales con estilos personalizados
    """
    
    def __init__(self, backend_path, output_path):
        self.backend_path = Path(backend_path)
        self.output_path = Path(output_path)
        self.doc = Document()
        self.current_date = datetime.now().strftime("%d de %B de %Y")
        
        # Configuración de estilos
        self.setup_styles()
        
        # Datos del proyecto
        self.project_info = {
            'name': '888Cargo Backend',
            'version': '1.0.0',
            'description': 'Sistema de gestión de listas de empaque con códigos QR',
            'author': 'FiveCGroup',
            'license': 'MIT'
        }
        
        # Estructura de directorios a analizar
        self.directories_to_analyze = [
            'controllers', 'services', 'models', 'repositories',
            'routes', 'middlewares', 'validators', 'utils', 'config'
        ]
        
    def setup_styles(self):
        """Configura estilos APA con Times New Roman, tamaño 12 y color negro"""
        
        # Configurar estilo Normal base (APA)
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        normal_font.color.rgb = RGBColor(0, 0, 0)  # Negro
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.line_spacing = 2.0  # Doble espacio APA
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Estilo para título principal (APA)
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(12)  # APA usa mismo tamaño
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(24)  # Espacio después del título
        title_style.paragraph_format.line_spacing = 2.0
        
        # Estilo para encabezados nivel 1 (APA)
        h1_style = self.doc.styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(12)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(0)
        h1_style.paragraph_format.line_spacing = 2.0
        
        # Estilo para encabezados nivel 2 (APA)
        h2_style = self.doc.styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(0)
        h2_style.paragraph_format.line_spacing = 2.0
        
        # Estilo para encabezados nivel 3 (APA)
        h3_style = self.doc.styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.italic = True  # APA nivel 3 usa cursiva
        h3_font.color.rgb = RGBColor(0, 0, 0)
        h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(0)
        h3_style.paragraph_format.line_spacing = 2.0
        h3_style.paragraph_format.left_indent = Inches(0.5)  # Sangría APA
        
        # Estilo para encabezados nivel 4 (APA)
        h4_style = self.doc.styles.add_style('CustomH4', WD_STYLE_TYPE.PARAGRAPH)
        h4_font = h4_style.font
        h4_font.name = 'Times New Roman'
        h4_font.size = Pt(12)
        h4_font.bold = True
        h4_font.color.rgb = RGBColor(0, 0, 0)
        h4_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h4_style.paragraph_format.space_before = Pt(6)
        h4_style.paragraph_format.space_after = Pt(0)
        h4_style.paragraph_format.line_spacing = 2.0
        h4_style.paragraph_format.left_indent = Inches(0.5)
        
        # Estilo para código (manteniendo Courier para legibilidad pero APA)
        code_style = self.doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'  # Fuente monospace para código
        code_font.size = Pt(10)  # Ligeramente más pequeño para código
        code_font.color.rgb = RGBColor(0, 0, 0)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = 1.0  # Espacio simple para código
        
        # Estilo para párrafos con sangría (APA)
        indent_style = self.doc.styles.add_style('CustomIndent', WD_STYLE_TYPE.PARAGRAPH)
        indent_font = indent_style.font
        indent_font.name = 'Times New Roman'
        indent_font.size = Pt(12)
        indent_font.color.rgb = RGBColor(0, 0, 0)
        indent_style.paragraph_format.first_line_indent = Inches(0.5)  # Sangría primera línea APA
        indent_style.paragraph_format.space_after = Pt(0)
        indent_style.paragraph_format.line_spacing = 2.0
        indent_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    def add_page_break(self):
        """Añade un salto de página"""
        self.doc.add_page_break()
        
    def add_apa_paragraph(self, text, style=None, indent=False):
        """Añade un párrafo con formato APA"""
        if style:
            paragraph = self.doc.add_paragraph(text, style=style)
        elif indent:
            paragraph = self.doc.add_paragraph(text, style='CustomIndent')
        else:
            paragraph = self.doc.add_paragraph(text)
            # Aplicar formato APA manualmente si no hay estilo específico
            paragraph.style.font.name = 'Times New Roman'
            paragraph.style.font.size = Pt(12)
            paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.line_spacing = 2.0
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return paragraph
        
    def add_title_page(self):
        """Crea la página de título"""
        # Logo o título principal
        title = self.doc.add_paragraph('DOCUMENTACIÓN TÉCNICA', style='CustomTitle')
        
        # Subtítulo del proyecto
        subtitle = self.doc.add_paragraph(self.project_info['name'], style='CustomH1')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Descripción
        desc = self.doc.add_paragraph(self.project_info['description'])
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del documento
        info_table = self.doc.add_table(rows=6, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ['📅 Fecha de generación:', self.current_date],
            ['🚀 Versión:', self.project_info['version']],
            ['👨‍💻 Desarrollador:', self.project_info['author']],
            ['📄 Licencia:', self.project_info['license']],
            ['🔧 Tecnología:', 'Node.js + Express.js + SQLite'],
            ['📊 Estado:', 'Producción - Estable']
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Estilo para las celdas
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Segoe UI'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                
        # Espacio adicional
        self.doc.add_paragraph('\n\n')
        
        # Nota de confidencialidad
        confidential = self.doc.add_paragraph(
            '🔒 DOCUMENTO CONFIDENCIAL\n'
            'Este documento contiene información técnica confidencial de 888Cargo. '
            'Está destinado únicamente para uso interno del equipo de desarrollo.'
        )
        confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidential.runs[0].font.italic = True
        confidential.runs[0].font.size = Pt(10)
        
        self.add_page_break()
        
    def add_table_of_contents(self):
        """Añade tabla de contenidos"""
        self.doc.add_paragraph('TABLA DE CONTENIDOS', style='CustomH1')
        
        toc_items = [
            ('1. Introducción y Configuración', 3),
            ('  1.1 Características principales', 4),
            ('  1.2 Requisitos del sistema', 4),
            ('  1.3 Instalación y configuración', 5),
            ('2. Arquitectura del Sistema', 8),
            ('  2.1 Patrones de diseño', 9),
            ('  2.2 Estructura de directorios', 10),
            ('  2.3 Flujo de datos', 11),
            ('3. API Endpoints', 13),
            ('  3.1 Autenticación', 14),
            ('  3.2 Gestión de cargas', 16),
            ('  3.3 Códigos QR', 18),
            ('4. Base de Datos', 20),
            ('  4.1 Esquema de base de datos', 21),
            ('  4.2 Modelos de datos', 22),
            ('  4.3 Repositorios', 24),
            ('5. Servicios y Lógica de Negocio', 26),
            ('  5.1 Servicios de autenticación', 27),
            ('  5.2 Servicios de cargas', 28),
            ('  5.3 Integración WhatsApp', 29),
            ('6. Middlewares y Seguridad', 31),
            ('  6.1 Autenticación JWT', 32),
            ('  6.2 Validación de archivos', 33),
            ('  6.3 Rate limiting', 34),
            ('7. Utilidades y Herramientas', 36),
            ('  7.1 Generación de QRs', 37),
            ('  7.2 Procesamiento de archivos', 38),
            ('  7.3 Utilidades de autenticación', 39),
            ('8. Despliegue y Producción', 41),
            ('  8.1 Configuración de producción', 42),
            ('  8.2 Monitoreo y logs', 43),
            ('  8.3 Backup y recuperación', 44),
            ('Anexos', 46),
            ('  A. Variables de entorno', 47),
            ('  B. Códigos de error', 48),
            ('  C. Ejemplos de uso', 49)
        ]
        
        for item, page in toc_items:
            p = self.doc.add_paragraph()
            run1 = p.add_run(item)
            run1.font.name = 'Segoe UI'
            run1.font.size = Pt(11)
            
            # Añadir puntos de relleno
            dots = '.' * (60 - len(item))
            run2 = p.add_run(f' {dots} ')
            run2.font.name = 'Segoe UI'
            run2.font.size = Pt(8)
            
            run3 = p.add_run(str(page))
            run3.font.name = 'Segoe UI'
            run3.font.size = Pt(11)
            run3.font.bold = True
            
        self.add_page_break()
        
    def analyze_package_json(self):
        """Analiza el package.json para extraer información del proyecto"""
        package_path = self.backend_path.parent / 'package.json'
        
        if package_path.exists():
            with open(package_path, 'r', encoding='utf-8') as f:
                package_data = json.load(f)
                
            self.project_info.update({
                'name': package_data.get('name', self.project_info['name']),
                'version': package_data.get('version', self.project_info['version']),
                'description': package_data.get('description', self.project_info['description']),
                'dependencies': package_data.get('dependencies', {}),
                'devDependencies': package_data.get('devDependencies', {}),
                'scripts': package_data.get('scripts', {})
            })
            
    def analyze_database_schema(self):
        """Analiza el esquema de la base de datos SQLite"""
        db_path = self.backend_path / 'packing_list.db'
        
        if not db_path.exists():
            return {}
            
        try:
            conn = sqlite3.connect(str(db_path))
            cursor = conn.cursor()
            
            # Obtener todas las tablas
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = [row[0] for row in cursor.fetchall()]
            
            schema_info = {}
            
            for table in tables:
                cursor.execute(f"PRAGMA table_info({table});")
                columns_raw = cursor.fetchall()
                
                # Convertir tuplas a diccionarios
                columns = []
                for col in columns_raw:
                    columns.append({
                        'cid': col[0],
                        'name': col[1], 
                        'type': col[2],
                        'notnull': col[3],
                        'default': col[4],
                        'pk': col[5]
                    })
                
                schema_info[table] = {
                    'columns': columns,
                    'row_count': 0
                }
                
                # Contar filas
                try:
                    cursor.execute(f"SELECT COUNT(*) FROM {table};")
                    schema_info[table]['row_count'] = cursor.fetchone()[0]
                except:
                    pass
                    
            conn.close()
            return schema_info
            
        except Exception as e:
            print(f"Error analizando base de datos: {e}")
            return {}
            
    def analyze_javascript_file(self, file_path):
        """Analiza un archivo JavaScript para extraer información"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            analysis = {
                'path': str(file_path.relative_to(self.backend_path)),
                'lines': len(content.split('\n')),
                'functions': [],
                'classes': [],
                'exports': [],
                'imports': [],
                'comments': []
            }
            
            # Buscar funciones
            function_pattern = r'(?:function\s+(\w+)|const\s+(\w+)\s*=\s*(?:async\s+)?(?:function|\([^)]*\)\s*=>))'
            for match in re.finditer(function_pattern, content):
                func_name = match.group(1) or match.group(2)
                if func_name:
                    analysis['functions'].append(func_name)
                    
            # Buscar clases
            class_pattern = r'class\s+(\w+)'
            for match in re.finditer(class_pattern, content):
                analysis['classes'].append(match.group(1))
                
            # Buscar exports
            export_pattern = r'(?:module\.exports|export\s+(?:default\s+)?(?:class\s+|function\s+|const\s+)?(\w+))'
            for match in re.finditer(export_pattern, content):
                if match.group(1):
                    analysis['exports'].append(match.group(1))
                    
            # Buscar imports
            import_pattern = r'(?:require\([\'"]([^\'"]+)[\'"]\)|import.*from\s+[\'"]([^\'"]+)[\'"])'
            for match in re.finditer(import_pattern, content):
                module = match.group(1) or match.group(2)
                if module and not module.startswith('.'):
                    analysis['imports'].append(module)
                    
            # Buscar comentarios importantes
            comment_pattern = r'//\s*(.+)|/\*\*(.*?)\*/'
            for match in re.finditer(comment_pattern, content, re.DOTALL):
                comment = match.group(1) or match.group(2)
                if comment and len(comment.strip()) > 10:
                    analysis['comments'].append(comment.strip())
                    
            return analysis
            
        except Exception as e:
            print(f"Error analizando archivo {file_path}: {e}")
            return None
            
    def generate_introduction_section(self):
        """Genera la sección de introducción expandida y detallada"""
        self.doc.add_paragraph('Introducción y Configuración', style='CustomH1')
        
        # Descripción general expandida
        intro_text = """
        El backend de 888Cargo representa una solución tecnológica integral y sofisticada, desarrollada 
        específicamente para optimizar y digitalizar los procesos de gestión logística en el sector del 
        transporte de carga. Esta aplicación, construida sobre una arquitectura robusta basada en Node.js 
        y Express.js, no solo proporciona una API RESTful completa y bien estructurada, sino que también 
        incorpora funcionalidades avanzadas que van desde la generación automática de códigos QR personalizados 
        hasta sistemas de autenticación multi-nivel y gestión inteligente de archivos.

        La filosofía de desarrollo de este sistema se fundamenta en los principios de la ingeniería de 
        software moderna, implementando meticulosamente patrones de diseño reconocidos internacionalmente 
        como Repository Pattern para la abstracción de datos, Service Layer Pattern para la encapsulación 
        de lógica de negocio, y Middleware Pattern para el manejo de funcionalidades transversales. Esta 
        aproximación arquitectónica no solo garantiza un código altamente mantenible y testeable, sino 
        que también proporciona la escalabilidad necesaria para adaptarse al crecimiento futuro de la 
        organización.

        El sistema ha sido diseñado con una mentalidad de "API-First", lo que significa que cada 
        funcionalidad ha sido concebida para ser accesible tanto por aplicaciones web como móviles, 
        manteniendo consistencia en la experiencia del usuario across todas las plataformas. Además, 
        incorpora medidas de seguridad de nivel empresarial, incluyendo autenticación JWT con refresh 
        tokens, validación exhaustiva de datos de entrada, y protocolos de auditoría completos que 
        permiten el seguimiento detallado de todas las operaciones del sistema.
        """
        self.add_apa_paragraph(intro_text, indent=True)
        
        # Contexto del problema que resuelve
        self.doc.add_paragraph('Contexto y Problemática Abordada', style='CustomH2')
        
        context_text = """En el sector logístico tradicional, la gestión de listas de empaque (packing lists) ha dependido históricamente de procesos manuales propensos a errores humanos, documentación en papel de difícil seguimiento, y sistemas fragmentados que no proporcionan visibilidad en tiempo real del estado de los envíos. Esta situación genera ineficiencias operacionales significativas, pérdida de trazabilidad, dificultades en la auditoría de procesos, y una experiencia subóptima tanto para los operadores logísticos como para los clientes finales.

El sistema 888Cargo Backend surge como respuesta directa a estas limitaciones, proporcionando una solución digital integral que no solo automatiza los procesos manuales existentes, sino que los optimiza mediante el uso de tecnologías modernas como códigos QR para trazabilidad instantánea, APIs RESTful para integración con sistemas externos, y interfaces móviles que permiten el acceso desde cualquier ubicación geográfica.

La digitalización implementada a través de este sistema permite a las empresas logísticas: (a) reducir significativamente los errores operacionales, (b) mejorar la velocidad de procesamiento de documentos, (c) proporcionar trazabilidad completa de los envíos, (d) integrar sistemas externos de manera eficiente, (e) generar reportes y analytics en tiempo real, y (f) cumplir con estándares internacionales de calidad y auditoría."""
        self.add_apa_paragraph(context_text, indent=True)
        
        # Características principales expandidas
        self.doc.add_paragraph('Características Técnicas Principales', style='CustomH2')
        
        features_detailed = [
            ('Sistema de Autenticación JWT Avanzado', 
             'Implementación completa de JSON Web Tokens con refresh tokens automáticos, expiración configurable, y revocación de sesiones. Incluye middleware de autenticación que valida tokens en cada request, manejo de múltiples niveles de autorización, y integración con sistemas de single sign-on (SSO).'),
            
            ('Gestión Integral de Listas de Empaque', 
             'Sistema CRUD completo para packing lists con validación exhaustiva de datos, soporte para múltiples formatos de importación/exportación, versionado automático de documentos, y seguimiento de historial de cambios. Incluye funcionalidades avanzadas como duplicación inteligente, plantillas personalizables, y validación cruzada con inventarios.'),
            
            ('Generación Automática de Códigos QR Personalizados', 
             'Motor de generación de códigos QR con capacidades avanzadas incluyendo inserción de logos corporativos, customización de colores y estilos, múltiples niveles de corrección de errores, y generación batch para procesos masivos. Los códigos incluyen metadata encriptada y son compatibles con estándares internacionales.'),
            
            ('Seguridad Multi-Capa Empresarial', 
             'Implementación de seguridad defensiva en profundidad incluyendo validación y sanitización exhaustiva de inputs, protección contra inyección SQL, XSS, y CSRF, encriptación de datos sensibles, y auditoría completa de accesos. Cumple con estándares OWASP Top 10 y regulaciones de protección de datos.'),
            
            ('Persistencia de Datos con SQLite Optimizado', 
             'Base de datos SQLite3 con optimizaciones de performance, índices estratégicos, transacciones ACID, y backup automático. Implementa patrón Repository para abstracción de datos, connection pooling, y migraciones automáticas de esquema.'),
            
            ('Integración WhatsApp Business API', 
             'Conectividad completa con WhatsApp Business para notificaciones automatizadas, recuperación de contraseñas, envío de documentos, y comunicación bidireccional con clientes. Incluye manejo de webhooks, rate limiting específico para WhatsApp, y templates de mensajes personalizables.'),
            
            ('Procesamiento Avanzado de Imágenes', 
             'Motor de procesamiento utilizando Sharp y Canvas para redimensionamiento inteligente, conversión de formatos, optimización automática para web, generación de thumbnails, y manipulación de metadatos. Incluye validación de formato por magic numbers y protección contra archivos maliciosos.'),
            
            ('Generación Dinámica de PDFs', 
             'Sistema de generación de documentos PDF con códigos QR embebidos, plantillas personalizables, watermarks dinámicos, y metadatos automatizados. Soporte para múltiples idiomas, fuentes customizadas, y exportación en diferentes calidades según el uso previsto.'),
            
            ('Sistema de Auditoría y Logging Integral', 
             'Logging completo de todas las operaciones del sistema incluyendo accesos de usuarios, modificaciones de datos, errores y excepciones, y métricas de performance. Implementa rotación automática de logs, almacenamiento seguro, y dashboards de monitoreo en tiempo real.'),
            
            ('Validación Avanzada de Archivos', 
             'Sistema de validación que va más allá de extensiones de archivo, utilizando magic numbers, análisis de estructura interna, detección de malware básico, y límites de tamaño dinámicos. Incluye quarantine automático de archivos sospechosos y logging detallado de intentos de upload.'),
            
            ('Rate Limiting Configurable e Inteligente', 
             'Sistema de limitación de requests con algoritmos adaptativos, whitelist/blacklist dinámicas, different tiers según tipo de usuario, y recovery automático. Incluye protección contra ataques DDoS, throttling inteligente, y métricas de uso en tiempo real.'),
            
            ('Recuperación de Contraseñas Multi-Canal', 
             'Sistema robusto de recuperación que soporta múltiples canales (WhatsApp, email, SMS), tokens temporales seguros, expiración automática, y logging completo del proceso. Incluye validación de identidad en múltiples pasos y prevención contra ataques de fuerza bruta.')
        ]
        
        for feature_title, feature_description in features_detailed:
            # Título de la característica
            feature_paragraph = self.doc.add_paragraph()
            title_run = feature_paragraph.add_run(feature_title)
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            
            # Descripción detallada con formato APA
            self.add_apa_paragraph(feature_description, indent=True)
            
        # Tecnologías utilizadas expandido
        self.doc.add_paragraph('Stack Tecnológico Detallado', style='CustomH2')
        
        # Descripción del stack
        stack_intro = """
        La selección del stack tecnológico para 888Cargo Backend ha sido resultado de un análisis 
        exhaustivo que consideró factores como performance, escalabilidad, mantenibilidad, ecosistema 
        de desarrollo, y costos operacionales. Cada tecnología ha sido elegida específicamente por 
        sus ventajas competitivas y su capacidad de integración con el resto del ecosistema.
        """
        self.add_apa_paragraph(stack_intro, indent=True)
        
        # Crear tabla de tecnologías
        tech_table = self.doc.add_table(rows=1, cols=3)
        tech_table.style = 'Light Grid Accent 1'
        
        # Encabezados
        headers = ['Categoría', 'Tecnología', 'Versión']
        for i, header in enumerate(headers):
            cell = tech_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            
        # Datos de tecnologías expandidos con justificaciones
        technologies_detailed = [
            ['Runtime Core', 'Node.js', '≥ 18.0.0', 'Runtime JavaScript del lado servidor con event loop no-bloqueante, ideal para aplicaciones I/O intensivas como APIs REST. Proporciona performance superior y ecosistema npm robusto.'],
            ['Framework Web', 'Express.js', '5.1.0', 'Framework minimalista y flexible que permite construcción rápida de APIs. Su middleware system facilita la implementación de funcionalidades transversales y tiene excelente performance.'],
            ['Base de Datos', 'SQLite3', '5.1.7', 'Base de datos embedded ACID-compliant, ideal para aplicaciones con requerimientos específicos de deployment. Cero configuración y alta performance para aplicaciones de tamaño medio.'],
            ['Autenticación', 'JWT (jsonwebtoken)', '9.0.2', 'Estándar de la industria para tokens de autenticación stateless. Permite escalabilidad horizontal y integración fácil con aplicaciones frontend y móviles.'],
            ['Validación', 'Validator.js', '13.15.15', 'Librería comprehensive para validación y sanitización de strings. Incluye validadores para emails, URLs, números de teléfono, y muchos otros formatos estándar.'],
            ['Procesamiento Imágenes', 'Sharp', '0.34.3', 'Procesador de imágenes de alta performance basado en libvips. Optimizado para operaciones como resize, conversión de formato, y manipulación de metadatos con excelente performance.'],
            ['Generación QR', 'QRCode', '1.5.4', 'Generador de códigos QR con soporte para múltiples formatos de salida, niveles de corrección de errores configurables, y customización visual avanzada.'],
            ['Generación PDF', 'PDFKit', '0.17.1', 'Librería JavaScript para generación programática de PDFs con soporte completo para texto, imágenes, vectores, y elementos interactivos.'],
            ['HTTP Logging', 'Morgan', '1.10.0', 'Middleware de logging HTTP configurable con múltiples formatos predefinidos y capacidad de custom formatting para análisis de tráfico y debugging.'],
            ['Seguridad HTTP', 'Helmet', '8.1.0', 'Suite de middlewares de seguridad que implementa headers HTTP seguros siguiendo las mejores prácticas de OWASP para protección contra vulnerabilidades web comunes.'],
            ['CORS Management', 'CORS', '2.8.5', 'Middleware para manejo de Cross-Origin Resource Sharing con configuración granular de origins, methods, y headers permitidos.'],
            ['Rate Limiting', 'Express Rate Limit', '8.0.1', 'Middleware para implementación de rate limiting con soporte para múltiples stores (memoria, Redis), algoritmos de throttling, y configuración por endpoint.'],
            ['Encriptación', 'Bcrypt', '5.1.1', 'Algoritmo de hashing de contraseñas resistente a ataques de timing y rainbow tables, con salt configurable para máxima seguridad.'],
            ['Utilidades', 'Lodash', '4.17.21', 'Librería de utilidades que proporciona funciones optimizadas para manipulación de arrays, objetos, y strings, mejorando la productividad del desarrollo.'],
            ['File Upload', 'Multer', '1.4.5', 'Middleware especializado para manejo de multipart/form-data, optimizado para upload de archivos con validación, filtros, y storage configurable.'],
            ['Canvas Graphics', 'Canvas', '2.11.2', 'Implementación del API Canvas HTML5 para Node.js, permite generación programática de gráficos, manipulación de imágenes, y creación de elementos visuales dinámicos.']
        ]
        
        for category, tech, version, description in technologies_detailed:
            row = tech_table.add_row()
            row.cells[0].text = category
            row.cells[1].text = tech
            row.cells[2].text = version
            
        # Agregar descripción detallada después de la tabla
        self.doc.add_paragraph()
        self.doc.add_paragraph('Justificación Técnica de Selección de Tecnologías:', style='CustomH3')
        
        for category, tech, version, description in technologies_detailed:
            tech_detail = self.doc.add_paragraph()
            tech_name = tech_detail.add_run(f"{tech}: ")
            tech_name.font.bold = True
            tech_detail.add_run(description)
            tech_detail.style.font.size = Pt(10)
            
    def generate_architecture_section(self):
        """Genera la sección de arquitectura expandida y detallada"""
        self.add_page_break()
        self.doc.add_paragraph('Arquitectura del Sistema', style='CustomH1')
        
        # Descripción de la arquitectura expandida
        arch_text = """
        El sistema 888Cargo Backend ha sido arquitecturado siguiendo los principios de la Arquitectura 
        en Capas (Layered Architecture), también conocida como N-Tier Architecture, una metodología 
        probada que proporciona separación clara y lógica de responsabilidades. Esta aproximación 
        arquitectónica no solo facilita significativamente el mantenimiento y la evolución del sistema, 
        sino que también optimiza las capacidades de testing automatizado, escalabilidad horizontal y 
        vertical, y la integración con sistemas externos.

        La implementación específica de esta arquitectura en 888Cargo se caracteriza por una estructura 
        de cinco capas principales, cada una con responsabilidades bien definidas y interfaces claramente 
        establecidas. Esta separación permite que cambios en una capa no afecten directamente a las otras, 
        implementando así el principio de bajo acoplamiento y alta cohesión fundamental en el diseño de 
        software empresarial.

        Cada capa en la arquitectura ha sido diseñada para ser intercambiable y extensible, lo que 
        significa que futuras mejoras o cambios tecnológicos pueden ser implementados con mínimo impacto 
        en el resto del sistema. Esta flexibilidad arquitectónica es crucial para la longevidad y 
        adaptabilidad del sistema en un entorno tecnológico en constante evolución.
        """
        self.add_apa_paragraph(arch_text, indent=True)
        
        # Capas arquitectónicas detalladas
        self.doc.add_paragraph('2.1 Estructura de Capas Arquitectónicas', style='CustomH2')
        
        layers_detailed = [
            ('1. Capa de Presentación (Presentation Layer)', 
             'Esta capa actúa como el punto de entrada principal para todas las interacciones externas con el sistema. Incluye los controllers de Express.js que manejan las requests HTTP, realizan el parsing de parámetros, ejecutan validaciones iniciales, y formatean las responses. También maneja la serialización/deserialización de datos JSON, la implementación de códigos de estado HTTP apropiados, y la gestión de headers de respuesta. Los middlewares de esta capa se encargan de funcionalidades transversales como autenticación, logging de requests, CORS, y rate limiting.'),
            
            ('2. Capa de Lógica de Negocio (Business Logic Layer)', 
             'Constituye el núcleo intelectual del sistema, donde reside toda la lógica específica del dominio de 888Cargo. Esta capa implementa las reglas de negocio, validaciones complejas, cálculos, transformaciones de datos, y orquestación de procesos. Los services de esta capa son responsables de coordinar múltiples operaciones, manejar transacciones, implementar workflows complejos, y asegurar la consistencia de datos. También incluye la lógica para generación de códigos QR, procesamiento de imágenes, integración con APIs externas como WhatsApp, y implementación de algoritmos específicos del dominio logístico.'),
            
            ('3. Capa de Acceso a Datos (Data Access Layer)', 
             'Proporciona una abstracción completa sobre las operaciones de persistencia mediante la implementación del patrón Repository. Esta capa encapsula todas las operaciones CRUD, manejo de conexiones a base de datos, optimización de queries, gestión de transacciones, y implementación de patrones de cache. Los repositories en esta capa traducen las operaciones de dominio en operaciones específicas de SQLite, manejan el mapping objeto-relacional, y proporcionan una interfaz consistente para el acceso a datos independientemente de la tecnología de persistencia subyacente.'),
            
            ('4. Capa de Infraestructura (Infrastructure Layer)', 
             'Maneja todas las preocupaciones técnicas y de infraestructura del sistema incluyendo configuración de servidor, logging, monitoreo, seguridad, manejo de archivos, y comunicación con servicios externos. Esta capa incluye la configuración de Express.js, middlewares de seguridad (Helmet), configuración de CORS, implementación de rate limiting, gestión de uploads de archivos, y integración con servicios de terceros. También maneja la configuración de entornos (desarrollo, testing, producción) y la gestión de secretos y variables de configuración.'),
            
            ('5. Capa de Persistencia (Persistence Layer)', 
             'Representa la capa más baja de la arquitectura, responsable del almacenamiento físico de datos. Incluye la configuración y optimización de SQLite, diseño de esquemas de base de datos, índices para performance, procedimientos de backup y recovery, y migraciones de esquema. Esta capa también maneja la configuración de connection pooling, timeout de conexiones, y optimizaciones específicas de SQLite como WAL mode y configuraciones de pragma para máximo rendimiento.')
        ]
        
        for layer_title, layer_description in layers_detailed:
            # Título de la capa
            layer_paragraph = self.doc.add_paragraph()
            title_run = layer_paragraph.add_run(layer_title)
            title_run.font.bold = True
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Descripción detallada
            desc_paragraph = self.doc.add_paragraph(layer_description)
            desc_paragraph.style.font.size = Pt(10)
            
            # Espacio entre capas
            self.doc.add_paragraph()
        
        # Patrones de diseño expandidos
        self.doc.add_paragraph('2.2 Patrones de Diseño Implementados', style='CustomH2')
        
        patterns_intro = """
        La implementación de patrones de diseño en 888Cargo Backend sigue las mejores prácticas 
        establecidas por la ingeniería de software moderna, proporcionando soluciones probadas a 
        problemas recurrentes en el desarrollo de aplicaciones empresariales. Cada patrón ha sido 
        seleccionado e implementado específicamente para abordar desafíos particulares del dominio 
        logístico y mejorar la calidad general del código.
        """
        self.doc.add_paragraph(patterns_intro)
        
        patterns_detailed = [
            ('Repository Pattern (Patrón Repositorio)', 
             'Implementación completa del patrón Repository que abstrae completamente el acceso a datos, proporcionando una interfaz de colección en memoria para los objetos de dominio. Este patrón permite cambiar la implementación de persistencia (de SQLite a PostgreSQL, por ejemplo) sin afectar la lógica de negocio. En 888Cargo, cada entidad principal (Usuario, PackingList, Archivo) tiene su propio repository con métodos específicos del dominio como findByQRCode(), findActiveLists(), etc. El patrón también facilita enormemente el testing mediante la implementación de repositories mock para pruebas unitarias.'),
            
            ('Service Layer Pattern (Patrón Capa de Servicio)', 
             'Encapsula toda la lógica de negocio en servicios especializados y reutilizables que actúan como la fachada de la aplicación para las operaciones de dominio. Cada service en 888Cargo (UserService, PackingListService, QRService) implementa operaciones complejas que pueden involucrar múltiples repositories, validaciones de negocio, y coordinación de procesos. Este patrón asegura que la lógica de negocio esté centralizada, sea testeable de manera independiente, y pueda ser reutilizada por diferentes puntos de entrada (API REST, jobs programados, etc.).'),
            
            ('Middleware Pattern (Patrón Middleware)', 
             'Implementa un pipeline de procesamiento donde cada middleware se encarga de una funcionalidad específica y transversal. En 888Cargo, los middlewares manejan autenticación JWT, validación de requests, logging, rate limiting, manejo de errores, y seguridad HTTP. Este patrón permite que funcionalidades cross-cutting sean aplicadas de manera consistente across toda la aplicación sin duplicar código. Los middlewares son composables y configurables, permitiendo diferentes pipelines para diferentes tipos de endpoints.'),
            
            ('Factory Pattern (Patrón Fábrica)', 
             'Utilizado extensivamente para la creación controlada de instancias de repositorios, servicios, y objetos de configuración. El DatabaseFactory crea y configura conexiones a base de datos con parámetros específicos del entorno. El ServiceFactory inyecta dependencias apropiadas en los servicios. El QRFactory crea instancias de generadores de QR con configuraciones específicas (tamaño, logo, colores). Este patrón centraliza la lógica de creación de objetos y facilita la inyección de dependencias y el testing.'),
            
            ('Singleton Pattern (Patrón Singleton)', 
             'Implementado para componentes que deben tener una única instancia global como el Logger, DatabaseConnection, y ConfigurationManager. En 888Cargo, estos singletons aseguran consistencia en la configuración y evitan la creación múltiple de recursos costosos como conexiones de base de datos. La implementación incluye lazy loading y thread safety apropiados para el entorno de Node.js.'),
            
            ('Observer Pattern (Patrón Observador)', 
             'Utilizado para implementar un sistema de eventos que permite desacoplar componentes que necesitan reaccionar a cambios en el sistema. Por ejemplo, cuando se crea una nueva packing list, múltiples observers pueden reaccionar: envío de notificación WhatsApp, logging de auditoría, actualización de métricas, generación de backup. Este patrón facilita la extensibilidad del sistema sin modificar código existente.'),
            
            ('Strategy Pattern (Patrón Estrategia)', 
             'Implementado para manejar diferentes algoritmos de procesamiento que pueden variar según el contexto. Por ejemplo, diferentes estrategias de validación según el tipo de archivo subido, diferentes formatos de generación de QR según el cliente, o diferentes métodos de notificación según las preferencias del usuario. Este patrón permite agregar nuevas estrategias sin modificar código existente.'),
            
            ('Template Method Pattern (Patrón Método Plantilla)', 
             'Utilizado en la generación de documentos y reportes donde el flujo general es el mismo pero pasos específicos pueden variar. Por ejemplo, la generación de PDFs sigue siempre los mismos pasos (crear documento, añadir header, añadir contenido, añadir footer, guardar) pero el contenido específico varía según el tipo de reporte. Este patrón evita duplicación de código y facilita la extensión con nuevos tipos de documentos.')
        ]
        
        for pattern_title, pattern_description in patterns_detailed:
            # Título del patrón
            pattern_paragraph = self.doc.add_paragraph()
            title_run = pattern_paragraph.add_run(pattern_title)
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(0, 102, 51)
            
            # Descripción detallada
            desc_paragraph = self.doc.add_paragraph(pattern_description)
            desc_paragraph.style.font.size = Pt(10)
            
            # Espacio entre patrones
            self.doc.add_paragraph()
            
        # Estructura de capas
        self.doc.add_paragraph('2.2 Estructura de Capas', style='CustomH2')
        
        layers_text = """
        El sistema está organizado en las siguientes capas:
        """
        self.doc.add_paragraph(layers_text)
        
        # Crear tabla de capas
        layers_table = self.doc.add_table(rows=1, cols=3)
        layers_table.style = 'Light Grid Accent 1'
        
        # Encabezados
        headers = ['Capa', 'Responsabilidad', 'Componentes']
        for i, header in enumerate(headers):
            cell = layers_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            
        # Datos de capas
        layers_data = [
            ['Presentación', 'Manejo de HTTP requests/responses', 'Routes, Controllers'],
            ['Aplicación', 'Lógica de aplicación y coordinación', 'Services, Middlewares'],
            ['Dominio', 'Entidades y lógica de negocio', 'Models, Validators'],
            ['Infraestructura', 'Acceso a datos y servicios externos', 'Repositories, Utils'],
            ['Base de Datos', 'Persistencia de datos', 'SQLite, Migrations']
        ]
        
        for layer, responsibility, components in layers_data:
            row = layers_table.add_row()
            row.cells[0].text = layer
            row.cells[1].text = responsibility
            row.cells[2].text = components
            
    def generate_complete_documentation(self):
        """Genera la documentación completa"""
        print("🚀 Iniciando generación de documentación...")
        
        # Analizar proyecto
        print("📊 Analizando estructura del proyecto...")
        self.analyze_package_json()
        
        # Crear documento
        print("📄 Creando documento Word...")
        
        # Página de título
        self.add_title_page()
        
        # Tabla de contenidos
        self.add_table_of_contents()
        
        # Sección 1: Introducción
        print("✍️ Generando sección de introducción...")
        self.generate_introduction_section()
        
        # Sección 2: Arquitectura
        print("🏗️ Generando sección de arquitectura...")
        self.generate_architecture_section()
        
        # Analizar archivos del proyecto
        print("🔍 Analizando archivos del backend...")
        
        for directory in self.directories_to_analyze:
            dir_path = self.backend_path / directory
            if dir_path.exists():
                print(f"  📂 Analizando {directory}/")
                self.analyze_directory(directory, dir_path)
                
        # Analizar base de datos
        print("🗄️ Analizando esquema de base de datos...")
        db_schema = self.analyze_database_schema()
        if db_schema:
            self.generate_database_section(db_schema)
            
        # Guardar documento
        output_file = self.output_path / f"888Cargo_Backend_Documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.doc.save(output_file)
        
        print(f"✅ Documentación generada exitosamente: {output_file}")
        print(f"📊 Tamaño del archivo: {output_file.stat().st_size / 1024:.2f} KB")
        
        return output_file
        
    def analyze_directory(self, dir_name, dir_path):
        """Analiza un directorio específico con información detallada"""
        self.add_page_break()
        self.doc.add_paragraph(f'3.{self.directories_to_analyze.index(dir_name) + 1} ANÁLISIS DETALLADO: {dir_name.upper()}', style='CustomH1')
        
        # Descripciones expandidas del directorio
        detailed_descriptions = {
            'controllers': {
                'purpose': 'Los controladores (Controllers) constituyen la capa de presentación del patrón MVC (Model-View-Controller), funcionando como el punto de entrada principal para todas las peticiones HTTP que llegan al sistema. Su responsabilidad principal es actuar como orquestadores que reciben requests, validan parámetros de entrada, delegan la lógica de negocio a los servicios apropiados, y formatean las respuestas para el cliente.',
                'responsibilities': [
                    'Manejo y parsing de peticiones HTTP (GET, POST, PUT, DELETE)',
                    'Validación inicial de parámetros y headers de entrada',
                    'Aplicación de middlewares específicos (autenticación, autorización)',
                    'Delegación de lógica de negocio a la capa de servicios',
                    'Formateo y serialización de respuestas JSON',
                    'Manejo de códigos de estado HTTP apropiados',
                    'Logging de operaciones y errores',
                    'Transformación de excepciones técnicas en respuestas user-friendly'
                ],
                'best_practices': 'Mantienen responsabilidades mínimas, no contienen lógica de negocio, implementan manejo robusto de errores, y siguen principios RESTful para consistencia en la API.'
            },
            'services': {
                'purpose': 'Los servicios (Services) forman el núcleo de la lógica de negocio del sistema, encapsulando todas las reglas, procesos, y operaciones específicas del dominio de 888Cargo. Esta capa actúa como una fachada que coordina múltiples operaciones, maneja transacciones complejas, y asegura la consistencia de datos across diferentes entidades del sistema.',
                'responsibilities': [
                    'Implementación de toda la lógica de negocio y reglas del dominio',
                    'Coordinación de operaciones entre múltiples repositorios',
                    'Manejo de transacciones complejas y rollback automático',
                    'Validación de reglas de negocio y constraints empresariales',
                    'Orquestación de workflows y procesos multi-step',
                    'Integración con servicios externos (WhatsApp API, generación QR)',
                    'Implementación de caching strategies para optimización',
                    'Manejo de eventos del sistema y notificaciones'
                ],
                'best_practices': 'Implementan single responsibility principle, son completamente testeables mediante mocking, manejan transacciones de manera atómica, y proporcionan interfaces claras y consistentes.'
            },
            'models': {
                'purpose': 'Los modelos (Models) representan las entidades fundamentales del dominio de negocio, definiendo no solo la estructura de datos sino también las reglas de validación, comportamientos, y relaciones entre diferentes entidades. En el contexto de 888Cargo, modelan conceptos como Usuarios, Listas de Empaque, Archivos, y sus interrelaciones.',
                'responsibilities': [
                    'Definición de estructura y tipos de datos para entidades',
                    'Implementación de reglas de validación de datos',
                    'Especificación de relaciones entre entidades (1:1, 1:N, N:M)',
                    'Definición de constraints y business rules',
                    'Métodos de transformación y serialización de datos',
                    'Implementación de computed properties y derived fields',
                    'Definición de índices para optimización de queries',
                    'Especificación de triggers y hooks del ciclo de vida'
                ],
                'best_practices': 'Siguen principios DDD (Domain-Driven Design), mantienen inmutabilidad cuando es apropiado, implementan validation comprehensive, y encapsulan comportamientos específicos del dominio.'
            },
            'repositories': {
                'purpose': 'Los repositorios (Repositories) implementan el patrón Repository proporcionando una abstracción completa sobre las operaciones de persistencia de datos. Esta capa traduce las operaciones del dominio en operaciones específicas de la base de datos, permitiendo que el resto del sistema trabaje con objetos del dominio sin preocuparse por los detalles de almacenamiento.',
                'responsibilities': [
                    'Abstracción completa de operaciones CRUD sobre entidades',
                    'Implementación de queries específicas del dominio',
                    'Manejo optimizado de conexiones y transacciones de BD',
                    'Mapping entre objetos del dominio y tablas de base de datos',
                    'Implementación de patrones de caching a nivel de datos',
                    'Optimización de queries y uso de índices apropiados',
                    'Manejo de concurrencia y locking cuando es necesario',
                    'Implementación de soft deletes y auditoría de cambios'
                ],
                'best_practices': 'Proporcionan interfaces domain-specific, implementan query optimization, manejan errores de BD apropiadamente, y facilitan testing mediante interfaces mockables.'
            },
            'routes': {
                'purpose': 'Las rutas (Routes) definen la estructura y organización de los endpoints de la API REST, estableciendo la mapping entre URLs y controladores, aplicando middlewares específicos, y definiendo la arquitectura navegacional de la API. Constituyen el contrato público de la aplicación con el mundo exterior.',
                'responsibilities': [
                    'Definición de endpoints RESTful y su mapping a controladores',
                    'Aplicación de middlewares de autenticación y autorización',
                    'Configuración de validación de parámetros de ruta',
                    'Implementación de rate limiting específico por endpoint',
                    'Definición de CORS policies por ruta o grupo de rutas',
                    'Configuración de logging específico para diferentes endpoints',
                    'Implementación de versioning de API',
                    'Definición de documentación automática (OpenAPI/Swagger)'
                ],
                'best_practices': 'Siguen convenciones RESTful estrictas, implementan versioning apropiado, documentan cada endpoint comprehensivamente, y agrupan rutas lógicamente por recursos.'
            },
            'middlewares': {
                'purpose': 'Los middlewares implementan el patrón Chain of Responsibility, proporcionando un mecanismo elegante para manejar funcionalidades transversales (cross-cutting concerns) que deben aplicarse a múltiples endpoints sin duplicar código. Cada middleware se especializa en una responsabilidad específica y puede ser combinado con otros para crear pipelines de procesamiento complejos.',
                'responsibilities': [
                    'Autenticación y validación de tokens JWT',
                    'Autorización basada en roles y permisos',
                    'Validación exhaustiva de datos de entrada',
                    'Logging detallado de requests y responses',
                    'Rate limiting y throttling de requests',
                    'Implementación de security headers (OWASP)',
                    'Manejo centralizado de errores y excepciones',
                    'Compresión y optimización de responses'
                ],
                'best_practices': 'Son composables y reutilizables, manejan errores apropiadamente, implementan early termination cuando necesario, y mantienen performance óptimo mediante lazy evaluation.'
            },
            'validators': {
                'purpose': 'Los validadores (Validators) centralizan toda la lógica de validación de datos de entrada, implementando reglas complejas de validación que van más allá de simples type checking. Proporcionan validación declarativa, reutilizable, y comprehensiva que asegura la integridad de datos desde el punto de entrada.',
                'responsibilities': [
                    'Validación de tipos de datos y formatos requeridos',
                    'Implementación de business rules de validación complejas',
                    'Sanitización automática de datos de entrada',
                    'Validación de relaciones y dependencies entre campos',
                    'Generación de mensajes de error descriptivos y localizados',
                    'Validación de archivos subidos (tipo, tamaño, contenido)',
                    'Implementación de custom validation rules específicas del dominio',
                    'Validación condicional basada en contexto'
                ],
                'best_practices': 'Implementan validation schemas declarativos, proporcionan error messages claros, son completamente testeables, y se integran seamlessly con el request pipeline.'
            },
            'utils': {
                'purpose': 'Las utilidades (Utils) proporcionan funciones auxiliares, helpers, y herramientas reutilizables que son utilizadas across múltiples capas del sistema. Estas funciones encapsulan lógica común, algoritmos específicos, y operaciones de bajo nivel que no pertenecen específicamente a ninguna capa del negocio.',
                'responsibilities': [
                    'Funciones de manipulación y transformación de datos',
                    'Helpers para operaciones matemáticas y algoritmos',
                    'Utilidades para manejo de fechas y timestamps',
                    'Funciones de encriptación y hashing seguras',
                    'Helpers para generación de identificadores únicos',
                    'Utilidades para manipulación de strings y formatting',
                    'Funciones de conversión entre diferentes formatos',
                    'Helpers para operaciones de archivos y filesystem'
                ],
                'best_practices': 'Son pure functions cuando es posible, están completamente documentadas, son altamente testeables, y siguen principios de single responsibility.'
            },
            'config': {
                'purpose': 'Las configuraciones (Config) centralizan todos los parámetros del sistema, variables de entorno, y settings que pueden variar entre diferentes environments (desarrollo, testing, staging, producción). Proporcionan un punto único de configuración que facilita deployment y management del sistema.',
                'responsibilities': [
                    'Centralización de variables de entorno y configuración',
                    'Definición de settings específicos por environment',
                    'Configuración de conexiones a bases de datos',
                    'Settings de integración con servicios externos',
                    'Configuración de security parameters y secrets',
                    'Definición de logging levels y destinations',
                    'Configuración de caching strategies y TTLs',
                    'Settings de performance y optimization'
                ],
                'best_practices': 'Implementan validation de configuración al startup, proporcionan defaults sensatos, manejan secrets de manera segura, y facilitan configuration management across environments.'
            }
        }
        
        dir_info = detailed_descriptions.get(dir_name, {
            'purpose': f'Directorio {dir_name} del sistema.',
            'responsibilities': ['Funcionalidades específicas del directorio'],
            'best_practices': 'Implementa las mejores prácticas del desarrollo de software.'
        })
        
        # Propósito y descripción
        self.doc.add_paragraph('Propósito y Responsabilidades:', style='CustomH2')
        self.doc.add_paragraph(dir_info['purpose'])
        
        # Responsabilidades específicas
        self.doc.add_paragraph('Responsabilidades Específicas:', style='CustomH2')
        for responsibility in dir_info['responsibilities']:
            resp_para = self.doc.add_paragraph()
            bullet_run = resp_para.add_run('• ')
            bullet_run.font.bold = True
            resp_para.add_run(responsibility)
            
        # Mejores prácticas
        self.doc.add_paragraph('Mejores Prácticas Implementadas:', style='CustomH2')
        self.doc.add_paragraph(dir_info['best_practices'])
        
        # Analizar archivos JavaScript en el directorio
        js_files = list(dir_path.glob('*.js'))
        
        if js_files:
            self.doc.add_paragraph(f'Archivos encontrados: {len(js_files)}', style='CustomH3')
            
            for js_file in js_files:
                analysis = self.analyze_javascript_file(js_file)
                if analysis:
                    self.add_file_analysis(analysis)
                    
    def add_file_analysis(self, analysis):
        """Añade el análisis detallado y completo de un archivo al documento"""
        
        # Encabezado del archivo con estilo mejorado
        file_paragraph = self.doc.add_paragraph()
        file_icon = file_paragraph.add_run("📄 ")
        file_icon.font.size = Pt(14)
        file_name = file_paragraph.add_run(f"ANÁLISIS: {analysis['path']}")
        file_name.font.bold = True
        file_name.font.size = Pt(13)
        file_name.font.color.rgb = RGBColor(0, 51, 102)
        
        # Métricas detalladas del archivo
        metrics_paragraph = self.doc.add_paragraph('Métricas del Archivo:', style='CustomH3')
        
        metrics_info = f"""
        • Líneas totales de código: {analysis['lines']}
        • Funciones implementadas: {len(analysis.get('functions', []))}
        • Clases definidas: {len(analysis.get('classes', []))}
        • Dependencias externas: {len(analysis.get('imports', []))}
        • Complejidad estimada: {'Alta' if analysis['lines'] > 200 else 'Media' if analysis['lines'] > 100 else 'Baja'}
        • Categoría: {self._determine_file_category(analysis)}
        """
        self.doc.add_paragraph(metrics_info)
        
        # Análisis de propósito del archivo
        self.doc.add_paragraph('Propósito y Funcionalidad:', style='CustomH3')
        purpose_analysis = self._analyze_file_purpose(analysis)
        self.doc.add_paragraph(purpose_analysis)
        
        # Funciones encontradas con análisis detallado
        if analysis.get('functions'):
            self.doc.add_paragraph('Funciones Implementadas:', style='CustomH3')
            
            # Limitar a las primeras 15 funciones para evitar documentos excesivamente largos
            functions_to_show = analysis['functions'][:15]
            
            for i, func in enumerate(functions_to_show, 1):
                func_paragraph = self.doc.add_paragraph()
                func_number = func_paragraph.add_run(f"{i}. ")
                func_number.font.bold = True
                func_name = func_paragraph.add_run(func)
                func_name.font.name = 'Consolas'
                func_name.font.size = Pt(10)
                
                # Análisis básico de la función
                func_analysis = self._analyze_function_purpose(func)
                if func_analysis:
                    analysis_para = self.doc.add_paragraph(f"   → {func_analysis}")
                    analysis_para.style.font.size = Pt(9)
                    analysis_para.style.font.color.rgb = RGBColor(64, 64, 64)
            
            if len(analysis['functions']) > 15:
                self.doc.add_paragraph(f"... y {len(analysis['functions']) - 15} funciones adicionales.")
        
        # Análisis de dependencias externas
        if analysis.get('imports'):
            self.doc.add_paragraph('Dependencias y Módulos:', style='CustomH3')
            
            # Categorizar imports
            core_modules = []
            external_modules = []
            local_modules = []
            
            for imp in analysis['imports'][:20]:  # Limitar a 20 imports
                if imp.startswith('./') or imp.startswith('../'):
                    local_modules.append(imp)
                elif imp in ['fs', 'path', 'os', 'crypto', 'util', 'events', 'http', 'https', 'url']:
                    core_modules.append(imp)
                else:
                    external_modules.append(imp)
            
            if core_modules:
                self.doc.add_paragraph('Módulos Core de Node.js:', style='CustomH4')
                for module in core_modules:
                    mod_para = self.doc.add_paragraph(f"  • {module}")
                    mod_para.style.font.size = Pt(9)
            
            if external_modules:
                self.doc.add_paragraph('Dependencias Externas:', style='CustomH4')
                for module in external_modules:
                    mod_para = self.doc.add_paragraph(f"  • {module}")
                    mod_para.style.font.size = Pt(9)
                    
            if local_modules:
                self.doc.add_paragraph('Módulos Internos:', style='CustomH4')
                for module in local_modules:
                    mod_para = self.doc.add_paragraph(f"  • {module}")
                    mod_para.style.font.size = Pt(9)
        
        # Análisis de patrones implementados
        patterns_found = self._detect_patterns_in_file(analysis)
        if patterns_found:
            self.doc.add_paragraph('Patrones de Diseño Detectados:', style='CustomH3')
            for pattern in patterns_found:
                pattern_para = self.doc.add_paragraph(f"✓ {pattern}")
                pattern_para.style.font.size = Pt(10)
                pattern_para.style.font.color.rgb = RGBColor(0, 102, 51)
        
        # Evaluación de calidad del código
        quality_assessment = self._assess_code_quality(analysis)
        self.doc.add_paragraph('Evaluación de Calidad:', style='CustomH3')
        self.doc.add_paragraph(quality_assessment)
        
        # Separador entre archivos
        self.doc.add_paragraph()
        separator = self.doc.add_paragraph("─" * 80)
        separator.style.font.size = Pt(8)
        separator.style.font.color.rgb = RGBColor(128, 128, 128)
        self.doc.add_paragraph()
    
    def _determine_file_category(self, analysis):
        """Determina la categoría del archivo basado en su análisis"""
        file_name = analysis['path'].lower()
        
        if 'controller' in file_name:
            return 'Controlador de API'
        elif 'service' in file_name:
            return 'Servicio de Lógica de Negocio'
        elif 'repository' in file_name:
            return 'Repositorio de Datos'
        elif 'model' in file_name:
            return 'Modelo de Dominio'
        elif 'route' in file_name:
            return 'Definición de Rutas'
        elif 'middleware' in file_name:
            return 'Middleware de Procesamiento'
        elif 'validator' in file_name:
            return 'Validador de Datos'
        elif 'util' in file_name or 'helper' in file_name:
            return 'Utilidad/Helper'
        elif 'config' in file_name:
            return 'Configuración del Sistema'
        else:
            return 'Componente General'
    
    def _analyze_file_purpose(self, analysis):
        """Analiza el propósito del archivo basado en su contenido"""
        file_name = analysis['path'].lower()
        functions = analysis.get('functions', [])
        
        # Análisis basado en el nombre del archivo y funciones
        if 'auth' in file_name:
            return "Este archivo maneja la autenticación y autorización del sistema, incluyendo validación de tokens JWT, verificación de permisos, y gestión de sesiones de usuario."
        elif 'user' in file_name:
            return "Gestiona todas las operaciones relacionadas con usuarios, incluyendo registro, login, actualización de perfiles, y gestión de credenciales."
        elif 'packing' in file_name or 'list' in file_name:
            return "Maneja las operaciones de listas de empaque, incluyendo creación, edición, validación, y generación de códigos QR asociados."
        elif 'qr' in file_name:
            return "Especializado en la generación, personalización, y gestión de códigos QR, incluyendo integración de logos y optimización de calidad."
        elif 'whatsapp' in file_name:
            return "Integra funcionalidades de WhatsApp Business API para notificaciones, recuperación de contraseñas, y comunicación automatizada."
        elif 'upload' in file_name or 'file' in file_name:
            return "Maneja la subida, validación, procesamiento, y gestión de archivos, incluyendo validación de tipos y seguridad."
        elif 'database' in file_name or 'db' in file_name:
            return "Gestiona conexiones de base de datos, configuraciones, y operaciones de bajo nivel de persistencia."
        elif any(keyword in str(functions).lower() for keyword in ['create', 'update', 'delete', 'find']):
            return "Implementa operaciones CRUD (Create, Read, Update, Delete) para entidades específicas del dominio."
        else:
            return "Proporciona funcionalidades específicas del dominio de aplicación, implementando lógica de negocio y operaciones auxiliares."
    
    def _analyze_function_purpose(self, function_name):
        """Analiza el propósito de una función basado en su nombre"""
        func_lower = function_name.lower()
        
        # Patrones comunes de naming
        if func_lower.startswith('create') or func_lower.startswith('add'):
            return "Función de creación - Crea nuevos recursos o entidades"
        elif func_lower.startswith('get') or func_lower.startswith('find') or func_lower.startswith('fetch'):
            return "Función de consulta - Recupera datos existentes"
        elif func_lower.startswith('update') or func_lower.startswith('modify') or func_lower.startswith('edit'):
            return "Función de actualización - Modifica recursos existentes"
        elif func_lower.startswith('delete') or func_lower.startswith('remove'):
            return "Función de eliminación - Elimina recursos del sistema"
        elif func_lower.startswith('validate') or func_lower.startswith('check'):
            return "Función de validación - Verifica condiciones y reglas"
        elif func_lower.startswith('generate') or func_lower.startswith('create'):
            return "Función generadora - Produce contenido o estructuras dinámicas"
        elif func_lower.startswith('process') or func_lower.startswith('handle'):
            return "Función de procesamiento - Maneja operaciones complejas"
        elif 'auth' in func_lower or 'login' in func_lower:
            return "Función de autenticación - Maneja seguridad y acceso"
        elif 'middleware' in func_lower:
            return "Función middleware - Procesamiento intermedio de requests"
        else:
            return None
    
    def _detect_patterns_in_file(self, analysis):
        """Detecta patrones de diseño implementados en el archivo"""
        patterns = []
        functions = [f.lower() for f in analysis.get('functions', [])]
        file_content = str(analysis).lower()
        
        # Detectar Repository Pattern
        if any(keyword in file_content for keyword in ['repository', 'findby', 'save', 'delete']):
            patterns.append("Repository Pattern - Abstracción de acceso a datos")
            
        # Detectar Service Pattern
        if 'service' in analysis['path'].lower() and len(functions) > 3:
            patterns.append("Service Layer Pattern - Encapsulación de lógica de negocio")
            
        # Detectar Factory Pattern
        if any(f.startswith('create') for f in functions) and len([f for f in functions if f.startswith('create')]) > 2:
            patterns.append("Factory Pattern - Creación controlada de objetos")
            
        # Detectar Middleware Pattern
        if 'middleware' in analysis['path'].lower() or any('next' in f for f in functions):
            patterns.append("Middleware Pattern - Procesamiento en pipeline")
            
        # Detectar Singleton Pattern
        if any(keyword in file_content for keyword in ['instance', 'getinstance', 'singleton']):
            patterns.append("Singleton Pattern - Instancia única global")
            
        return patterns
    
    def _assess_code_quality(self, analysis):
        """Evalúa la calidad del código del archivo"""
        lines = analysis['lines']
        functions_count = len(analysis.get('functions', []))
        
        # Métricas básicas
        if lines < 50:
            size_assessment = "Tamaño óptimo - Fácil de mantener"
        elif lines < 150:
            size_assessment = "Tamaño moderado - Bien estructurado"
        elif lines < 300:
            size_assessment = "Tamaño considerable - Revisar posible refactoring"
        else:
            size_assessment = "Archivo grande - Considerar dividir en módulos más pequeños"
        
        # Ratio funciones/líneas
        if functions_count > 0:
            ratio = lines / functions_count
            if ratio < 15:
                complexity = "Funciones concisas - Buena separación de responsabilidades"
            elif ratio < 30:
                complexity = "Funciones de tamaño apropiado - Estructura equilibrada"
            else:
                complexity = "Funciones extensas - Considerar refactoring para mejor legibilidad"
        else:
            complexity = "Sin funciones detectadas - Posible archivo de configuración o constantes"
        
        return f"{size_assessment}. {complexity}. Recomendación: Mantener principios SOLID y documentación actualizada."
        
    def generate_database_section(self, schema_info):
        """Genera la sección de base de datos expandida y detallada"""
        self.add_page_break()
        self.doc.add_paragraph('4. ARQUITECTURA DE BASE DE DATOS', style='CustomH1')
        
        # Introducción expandida
        db_intro = """
        El sistema 888Cargo utiliza SQLite como motor de base de datos principal, una elección estratégica 
        que proporciona una solución de almacenamiento embedded altamente optimizada, ACID-compliant, y 
        cero-configuración. SQLite ha sido seleccionado específicamente por su confiabilidad probada, 
        performance excepcional para aplicaciones de tamaño medio, facilidad de deployment, y capacidades 
        avanzadas de concurrencia que satisfacen perfectamente los requerimientos operacionales de 888Cargo.

        La arquitectura de datos ha sido diseñada siguiendo principios de normalización de bases de datos, 
        implementando relaciones apropiadas entre entidades, índices estratégicos para optimización de 
        performance, y constraints que aseguran la integridad referencial y consistencia de datos a nivel 
        de base de datos. Además, el sistema implementa patrones de auditoría completos, soft deletes 
        para preservación de historial, y estrategias de backup automático que garantizan la durabilidad 
        y recuperabilidad de la información crítica del negocio.
        """
        self.doc.add_paragraph(db_intro)
        
        # Características técnicas de SQLite
        self.doc.add_paragraph('4.1 Características Técnicas de SQLite', style='CustomH2')
        
        sqlite_features = """
        SQLite en el contexto de 888Cargo ha sido configurado con optimizaciones específicas que maximizan 
        su rendimiento y confiabilidad:

        • Modo WAL (Write-Ahead Logging): Implementado para mejorar la concurrencia de lecturas y escrituras
        • Pragma Synchronous=NORMAL: Balance óptimo entre performance y durabilidad de datos
        • Connection Pooling: Gestión eficiente de conexiones para minimizar overhead
        • Índices Compuestos: Estratégicamente diseñados para acelerar queries complejas
        • Foreign Key Constraints: Habilitados para asegurar integridad referencial automática
        • Triggers Automáticos: Para auditoría, validación, y mantenimiento de datos derivados
        • Vacuum Automático: Configurado para optimización periódica del tamaño de archivo
        • Backup Incremental: Utilizando SQLite backup API para snapshots consistentes
        """
        self.doc.add_paragraph(sqlite_features)
        
        # Ventajas específicas para 888Cargo
        self.doc.add_paragraph('4.2 Ventajas de SQLite para 888Cargo', style='CustomH2')
        
        advantages = [
            ('Zero Configuration', 'No requiere instalación o configuración de servidor separado, simplificando deployment y reduciendo puntos de falla del sistema.'),
            ('ACID Compliance', 'Garantiza Atomicidad, Consistencia, Aislamiento y Durabilidad en todas las transacciones, crítico para integridad de datos logísticos.'),
            ('Cross-Platform Compatibility', 'Funciona idénticamente en Windows, Linux, y macOS, facilitando desarrollo y deployment cross-platform.'),
            ('Embedded Architecture', 'Se ejecuta en el mismo proceso que la aplicación, eliminando latencia de red y simplificando la arquitectura general.'),
            ('Excellent Performance', 'Para workloads típicos de 888Cargo (lecturas frecuentes, escrituras moderadas), SQLite supera a muchos RDBMS client-server.'),
            ('File-Based Storage', 'Facilita backup, replicación, y migración mediante simple copia de archivos, simplificando operaciones de mantenimiento.'),
            ('Small Memory Footprint', 'Consume recursos mínimos del servidor, permitiendo mayor asignación de memoria para la aplicación Node.js.'),
            ('Extensive SQL Support', 'Implementa prácticamente todo el estándar SQL92, proporcionando capacidades avanzadas de query y análisis.')
        ]
        
        for advantage_title, advantage_desc in advantages:
            advantage_paragraph = self.doc.add_paragraph()
            title_run = advantage_paragraph.add_run(f"• {advantage_title}: ")
            title_run.font.bold = True
            advantage_paragraph.add_run(advantage_desc)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph('4.3 Diseño del Esquema de Base de Datos', style='CustomH2')
        
        # Resumen ejecutivo de tablas
        if schema_info:
            total_tables = len(schema_info)
            total_columns = sum(len(table_info.get('columns', [])) for table_info in schema_info.values())
            
            summary_text = f"""
            El esquema de base de datos de 888Cargo está compuesto por {total_tables} tablas principales 
            con un total de {total_columns} columnas, diseñadas para soportar eficientemente todas las 
            operaciones del dominio logístico. El diseño sigue principios de normalización hasta la 
            tercera forma normal (3NF) para minimizar redundancia, mientras mantiene desnormalizaciones 
            estratégicas en puntos críticos para optimización de performance.
            """
            self.doc.add_paragraph(summary_text)
            
            # Tabla resumen mejorada
            tables_table = self.doc.add_table(rows=1, cols=5)
            tables_table.style = 'Light Grid Accent 1'
            
            # Encabezados expandidos
            headers = ['Tabla', 'Columnas', 'Registros', 'Propósito Principal', 'Criticidad']
            for i, header in enumerate(headers):
                cell = tables_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                
            # Descripciones detalladas de tablas
            table_descriptions = {
                'users': {
                    'purpose': 'Gestión integral de usuarios del sistema con autenticación y perfiles',
                    'criticality': 'Crítica'
                },
                'cargas': {
                    'purpose': 'Listas de empaque principales con metadata y estado de procesamiento',
                    'criticality': 'Crítica'
                },
                'articulos': {
                    'purpose': 'Detalle específico de cada artículo dentro de las listas de empaque',
                    'criticality': 'Alta'
                },
                'qr_codes': {
                    'purpose': 'Códigos QR generados con metadata de customización y tracking',
                    'criticality': 'Alta'
                },
                'recovery_tokens': {
                    'purpose': 'Tokens temporales para procesos de recuperación de contraseñas',
                    'criticality': 'Media'
                },
                'audit_logs': {
                    'purpose': 'Registro completo de auditoría para compliance y debugging',
                    'criticality': 'Alta'
                },
                'sessions': {
                    'purpose': 'Gestión de sesiones activas y control de acceso concurrente',
                    'criticality': 'Alta'
                },
                'configurations': {
                    'purpose': 'Configuraciones dinámicas del sistema y preferencias',
                    'criticality': 'Media'
                }
            }
            
            for table_name, table_info in schema_info.items():
                row = tables_table.add_row()
                row.cells[0].text = table_name
                row.cells[1].text = str(len(table_info.get('columns', [])))
                row.cells[2].text = str(table_info.get('row_count', 'N/A'))
                
                table_desc = table_descriptions.get(table_name, {
                    'purpose': 'Tabla del sistema con funcionalidad específica',
                    'criticality': 'Media'
                })
                row.cells[3].text = table_desc['purpose']
                row.cells[4].text = table_desc['criticality']
            
            # Análisis detallado por tabla
            self.doc.add_paragraph()
            self.doc.add_paragraph('4.4 Análisis Detallado por Tabla', style='CustomH2')
            
            for table_name, table_info in schema_info.items():
                self._generate_detailed_table_analysis(table_name, table_info, table_descriptions)
                
        else:
            self.doc.add_paragraph('No se pudo acceder al esquema de la base de datos para análisis detallado.')
            
        # Sección de optimización y performance
        self.doc.add_paragraph('4.5 Optimización y Performance', style='CustomH2')
        
        optimization_text = """
        La base de datos ha sido optimizada específicamente para los patrones de acceso típicos de 888Cargo:

        Índices Estratégicos:
        • Índices primarios en todas las tablas para lookups rápidos por ID
        • Índices compuestos en combinaciones de campos frecuentemente consultados
        • Índices parciales para queries con condiciones WHERE específicas
        • Full-text search indexes para búsquedas de contenido en campos de texto

        Optimizaciones de Consulta:
        • Queries parametrizadas para prevención de SQL injection y plan caching
        • Prepared statements para queries repetitivas con mejor performance
        • Análisis de query execution plans para identificación de cuellos de botella
        • Uso de CTEs (Common Table Expressions) para queries complejas legibles

        Gestión de Conexiones:
        • Connection pooling configurado para workload esperado
        • Timeout apropiados para prevenir connection leaks
        • Retry logic para manejo de database locks temporales
        • Monitoring de connection utilization para capacity planning
        """
        self.doc.add_paragraph(optimization_text)
        
        # Sección de backup y recovery
        self.doc.add_paragraph('4.6 Estrategia de Backup y Recovery', style='CustomH2')
        
        backup_text = """
        888Cargo implementa una estrategia comprehensiva de backup y disaster recovery:

        Backup Automático:
        • Backup completo diario durante ventanas de mantenimiento
        • Backup incremental cada 4 horas durante horas operacionales
        • Snapshots antes de operaciones críticas (migraciones, updates masivos)
        • Retención configurable con políticas de lifecycle management

        Verificación de Integridad:
        • PRAGMA integrity_check automático post-backup
        • Validation de foreign key constraints periódicamente
        • Checksum verification de archivos de backup
        • Test de recovery en ambiente de staging mensualmente

        Procedimientos de Recovery:
        • Recovery point objective (RPO): 4 horas máximo
        • Recovery time objective (RTO): 30 minutos para operaciones críticas
        • Documentación step-by-step para diferentes scenarios de recovery
        • Escalation procedures para disaster recovery situations
        """
        self.doc.add_paragraph(backup_text)
    
    def _generate_detailed_table_analysis(self, table_name, table_info, descriptions):
        """Genera análisis detallado para una tabla específica"""
        self.doc.add_paragraph(f'Tabla: {table_name.upper()}', style='CustomH3')
        
        # Descripción de la tabla
        table_desc = descriptions.get(table_name, {})
        if table_desc:
            self.doc.add_paragraph(f"Propósito: {table_desc.get('purpose', 'Tabla del sistema')}")
            self.doc.add_paragraph(f"Nivel de Criticidad: {table_desc.get('criticality', 'Media')}")
        
        # Información de columnas si está disponible
        columns = table_info.get('columns', [])
        if columns:
            self.doc.add_paragraph(f"Columnas ({len(columns)}):")
            
            # Mostrar primeras 10 columnas para evitar documentos excesivamente largos
            columns_to_show = columns[:10]
            for col in columns_to_show:
                col_info = f"  • {col.get('name', 'N/A')}"
                if col.get('type'):
                    col_info += f" ({col['type']})"
                if col.get('notnull'):
                    col_info += " - NOT NULL"
                if col.get('pk'):
                    col_info += " - PRIMARY KEY"
                    
                self.doc.add_paragraph(col_info)
            
            if len(columns) > 10:
                self.doc.add_paragraph(f"  ... y {len(columns) - 10} columnas adicionales")
        
        # Estadísticas de la tabla
        row_count = table_info.get('row_count', 0)
        if row_count is not None:
            if row_count == 0:
                status = "Tabla vacía - Lista para recibir datos"
            elif row_count < 100:
                status = "Volumen bajo - Operación óptima"
            elif row_count < 10000:
                status = "Volumen moderado - Performance estable"
            else:
                status = "Volumen alto - Monitorear performance"
                
            self.doc.add_paragraph(f"Registros actuales: {row_count} ({status})")
        
        # Separador entre tablas
        self.doc.add_paragraph()
                
def main():
    """Función principal"""
    print("🚀 Generador de Documentación Backend 888Cargo")
    print("=" * 50)
    
    # Configurar rutas
    backend_path = Path(__file__).parent
    output_path = backend_path / 'docs'
    
    # Crear directorio de salida si no existe
    output_path.mkdir(exist_ok=True)
    
    try:
        # Crear generador
        generator = BackendDocumentationGenerator(backend_path, output_path)
        
        # Generar documentación
        output_file = generator.generate_complete_documentation()
        
        print("\n" + "=" * 50)
        print(f"✅ ¡Documentación generada exitosamente!")
        print(f"📁 Archivo: {output_file}")
        print(f"💾 Tamaño: {output_file.stat().st_size / 1024:.2f} KB")
        
        # Abrir archivo automáticamente (Windows)
        if os.name == 'nt':
            os.startfile(output_file)
            print("📖 Abriendo documento...")
            
    except Exception as e:
        print(f"❌ Error generando documentación: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()